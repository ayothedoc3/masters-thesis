"""
Generate front matter components for the thesis:
- Flyleaf
- Bilingual abstract
- Acknowledgments
- List of abbreviations
- Introduction

These documents are assembled into the final thesis by scripts/assemble_thesis.py.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CHAPTERS = os.path.join(ROOT, "chapters")
os.makedirs(CHAPTERS, exist_ok=True)

TITLE_EN = (
    "QUALITY OF PHYSICAL ACTIVITY INFORMATION ON TRADITIONAL WEBSITES "
    "VERSUS SOCIAL MEDIA PLATFORMS: A CROSS-SECTIONAL CONTENT ANALYSIS "
    "USING DISCERN AND JAMA BENCHMARKS"
)
TITLE_LT = (
    "FIZINIO AKTYVUMO INFORMACIJOS KOKYBĖ TRADICINĖSE INTERNETO "
    "SVETAINĖSE IR SOCIALINIŲ TINKLŲ PLATFORMOSE: SKERSPJŪVIO TURINIO "
    "ANALIZĖ TAIKANT DISCERN IR JAMA KRITERIJUS"
)


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    size: Pt = Pt(12),
    spacing_after: Pt | None = None,
    first_line_indent_cm: float | None = None,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = size
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 1.5
    if spacing_after is not None:
        p.paragraph_format.space_after = spacing_after
    if first_line_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)


def add_front_heading(doc: Document, text: str) -> None:
    add_para(
        doc,
        text,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        size=Pt(14),
        spacing_after=Pt(12),
    )


def add_rule_line(doc: Document, label: str) -> None:
    add_para(doc, f"{label} " + "_" * 35, align=WD_ALIGN_PARAGRAPH.LEFT)


def generate_flyleaf() -> None:
    doc = Document()
    set_style(doc)

    add_front_heading(doc, "CONFIRMATION OF INDEPENDENT COMPOSITION OF THE THESIS")
    add_para(
        doc,
        f"I hereby declare that the present final master's thesis titled "
        f"\"{TITLE_EN}\"",
        first_line_indent_cm=1.25,
    )
    for item in [
        "1. Has been carried out by myself;",
        "2. Has not been used in any other university in Lithuania or abroad;",
        "3. Has not used any references not indicated in the paper and the list of references is complete.",
    ]:
        add_para(doc, item, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_para(doc, "")
    add_rule_line(doc, "Date:")
    add_rule_line(doc, "Author's name and surname:")
    add_rule_line(doc, "Signature:")

    add_para(doc, "")
    add_front_heading(doc, "CONFIRMATION OF LIABILITY FOR THE REGULARITY OF THE ENGLISH LANGUAGE")
    add_para(
        doc,
        "I hereby confirm the correctness of the English language used in the final thesis.",
        first_line_indent_cm=1.25,
    )
    add_para(doc, "")
    add_rule_line(doc, "Date:")
    add_rule_line(doc, "Author's name and surname:")
    add_rule_line(doc, "Signature:")

    add_para(doc, "")
    add_front_heading(doc, "FINAL MASTER'S THESIS SUPERVISOR'S ASSESSMENT")
    add_rule_line(doc, "Date:")
    add_rule_line(doc, "Supervisor's name and surname:")
    add_rule_line(doc, "Signature:")
    add_para(doc, "")
    add_rule_line(doc, "Reviewer 1:")
    add_rule_line(doc, "Study administrator:")
    add_rule_line(doc, "Signature:")
    add_para(doc, "")
    add_rule_line(doc, "Reviewer 2:")
    add_rule_line(doc, "Study administrator:")
    add_rule_line(doc, "Signature:")
    add_para(doc, "")
    add_rule_line(doc, "Final thesis placed in ETD IS by:")
    add_rule_line(doc, "Date / signature:")

    path = os.path.join(CHAPTERS, "flyleaf.docx")
    doc.save(path)
    print(f"Flyleaf saved to {path}")


def generate_abstract() -> None:
    doc = Document()
    set_style(doc)

    add_front_heading(doc, "ABSTRACT")
    add_para(doc, TITLE_EN, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, TITLE_LT, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Keywords: physical activity; health information quality; social media; DISCERN; JAMA benchmarks",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        spacing_after=Pt(12),
    )
    add_para(
        doc,
        "Problem. The public increasingly relies on websites and social media for physical activity advice, "
        "yet evidence comparing the quality of information across these environments remains limited. "
        "Aim. To compare the quality of physical activity information on traditional websites, YouTube, "
        "TikTok, and Instagram using DISCERN and JAMA benchmarks. Objectives. To examine platform "
        "differences, creator-type differences, and associations between engagement metrics and quality. "
        "Hypothesis. Website content and content produced by healthcare or certified fitness professionals "
        "would score higher than short-form social media content, whereas engagement metrics would not "
        "reliably indicate quality. Methods. A cross-sectional content analysis was conducted on 197 "
        "publicly accessible English-language items retrieved with five standardized search terms. "
        "Information quality was assessed with the 16-item DISCERN instrument and four JAMA criteria. "
        "Inter-rater reliability was evaluated on 40 items. Kruskal-Wallis, Mann-Whitney U, chi-square, "
        "and Spearman analyses were used. Main findings. Quality differed significantly by platform: "
        "websites ranked highest, followed by YouTube, TikTok, and Instagram. JAMA compliance was low "
        "across all platforms. Creator type was associated with DISCERN scores, with healthcare and "
        "certified fitness professionals performing better than general users. On YouTube, higher views "
        "and likes were associated with lower quality. Conclusions. Traditional websites remain the most "
        "reliable source of physical activity information, whereas short-form social media content is more "
        "vulnerable to low quality and weak transparency.",
        first_line_indent_cm=1.25,
    )

    doc.add_page_break()

    add_front_heading(doc, "SANTRAUKA")
    add_para(doc, TITLE_LT, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, TITLE_EN, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "Raktiniai žodžiai: fizinis aktyvumas; sveikatos informacijos kokybė; socialiniai tinklai; DISCERN; JAMA kriterijai",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        spacing_after=Pt(12),
    )
    add_para(
        doc,
        "Problema. Visuomenė vis dažniau fizinio aktyvumo informacijos ieško interneto svetainėse ir "
        "socialiniuose tinkluose, tačiau lyginamųjų duomenų apie informacijos kokybę skirtingose "
        "skaitmeninėse aplinkose vis dar trūksta. Tikslas. Palyginti fizinio aktyvumo informacijos kokybę "
        "tradicinėse interneto svetainėse, YouTube, TikTok ir Instagram, taikant DISCERN instrumentą ir "
        "JAMA kriterijus. Uždaviniai. Įvertinti platformų skirtumus, kūrėjų tipų skirtumus ir ryšį tarp "
        "įsitraukimo rodiklių bei informacijos kokybės. Hipotezė. Manyta, kad tradicinių interneto "
        "svetainių turinys ir sveikatos priežiūros ar sertifikuotų fitneso specialistų sukurtas turinys "
        "bus aukštesnės kokybės nei trumpo formato socialinių tinklų turinys, o įsitraukimo rodikliai "
        "patikimai neatspindės kokybės. Metodai. Atlikta 197 viešai prieinamų anglų kalba paskelbtų "
        "vienetų skerspjūvio turinio analizė, naudojant penkias standartizuotas paieškos frazes. "
        "Informacijos kokybė vertinta 16 klausimų DISCERN instrumentu ir keturiais JAMA kriterijais. "
        "Tarpvertintojų patikimumas vertintas 40 vienetų imtyje. Taikyti Kruskal-Wallis, Mann-Whitney U, "
        "chi kvadrato ir Spearman koreliacijos testai. Pagrindiniai rezultatai. Informacijos kokybė "
        "statistiškai reikšmingai skyrėsi tarp platformų: aukščiausius rodiklius pasiekė interneto "
        "svetainės, po jų sekė YouTube, TikTok ir Instagram. JAMA kriterijų atitikimas visose platformose "
        "buvo žemas. Kūrėjo tipas buvo susijęs su DISCERN įverčiais: geriau pasirodė sveikatos priežiūros "
        "ir sertifikuoti fizinio aktyvumo specialistai nei bendrieji naudotojai. YouTube platformoje "
        "didesnis peržiūrų ir patiktukų skaičius buvo susijęs su prastesne kokybe. Išvados. Tradicinės "
        "interneto svetainės išlieka patikimiausiu fizinio aktyvumo informacijos šaltiniu, o trumpo "
        "formato socialinių tinklų turinys dažniau pasižymi prastesne kokybe ir mažesniu skaidrumu.",
        first_line_indent_cm=1.25,
    )

    path = os.path.join(CHAPTERS, "abstract.docx")
    doc.save(path)
    print(f"Abstract saved to {path}")


def generate_acknowledgments() -> None:
    doc = Document()
    set_style(doc)

    add_front_heading(doc, "ACKNOWLEDGMENTS")
    for paragraph in [
        (
            "I would like to express my sincere gratitude to my supervisor, Dr. Antanas Usas, "
            "for his guidance, expertise, and constructive feedback throughout this research project."
        ),
        (
            "I am grateful to the faculty and staff of the Master of Public Health programme at "
            "Lithuanian Sports University for providing the academic environment and resources "
            "necessary to complete this thesis."
        ),
        (
            "I also thank my family and friends for their patience and support during the course "
            "of my studies."
        ),
    ]:
        add_para(doc, paragraph, first_line_indent_cm=1.25)

    path = os.path.join(CHAPTERS, "acknowledgments.docx")
    doc.save(path)
    print(f"Acknowledgments saved to {path}")


def generate_abbreviations() -> None:
    doc = Document()
    set_style(doc)

    add_front_heading(doc, "LIST OF ABBREVIATIONS")

    abbreviations = [
        ("APA", "American Psychological Association"),
        ("CI", "Confidence Interval"),
        ("DISCERN", "Quality criteria for consumer health information instrument"),
        ("ICC", "Intraclass Correlation Coefficient"),
        ("IG", "Instagram"),
        ("JAMA", "Journal of the American Medical Association"),
        ("Mdn", "Median"),
        ("MPH", "Master of Public Health"),
        ("PA", "Physical Activity"),
        ("RQ", "Research Question"),
        ("TT", "TikTok"),
        ("WEB", "Traditional Website"),
        ("WHO", "World Health Organization"),
        ("YT", "YouTube"),
    ]

    table = doc.add_table(rows=len(abbreviations) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Abbreviation"
    table.rows[0].cells[1].text = "Full Term"
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True

    for i, (abbr, full) in enumerate(abbreviations, start=1):
        table.rows[i].cells[0].text = abbr
        table.rows[i].cells[1].text = full
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)

    path = os.path.join(CHAPTERS, "abbreviations.docx")
    doc.save(path)
    print(f"Abbreviations saved to {path}")


def generate_introduction() -> None:
    doc = Document()
    set_style(doc)

    heading = doc.add_heading("INTRODUCTION", level=1)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    add_para(
        doc,
        "Digital media have become a primary route through which the public seeks exercise and "
        "physical activity advice. Traditional health websites now coexist with algorithm-driven "
        "social media platforms that reward brevity, visual appeal, and engagement. Although this "
        "shift expands access to health information, it also creates uncertainty about whether the "
        "content most visible to users is accurate, transparent, and useful for health decision-making.",
        first_line_indent_cm=1.25,
    )
    add_para(
        doc,
        "The research problem addressed in this thesis is the uneven and insufficiently compared "
        "quality of physical activity information across major online platforms. Previous studies "
        "have usually examined single platforms, single health topics, or single assessment tools, "
        "which limits direct cross-platform interpretation. As a result, public health practitioners "
        "and consumers still lack clear evidence on whether websites, YouTube, TikTok, and Instagram "
        "differ in the quality and accountability of physical activity information delivered to "
        "ordinary users.",
        first_line_indent_cm=1.25,
    )
    add_para(
        doc,
        "The aim of this thesis is to compare the quality of physical activity information on "
        "traditional websites and three major social media platforms using DISCERN and JAMA benchmarks. "
        "The objectives are: to assess platform differences in information quality; to examine whether "
        "creator type is associated with quality; and to test whether engagement indicators are related "
        "to quality. The working hypothesis was that traditional websites and professionally produced "
        "content would score higher than short-form social media content, whereas engagement metrics "
        "would not reliably signal information quality.",
        first_line_indent_cm=1.25,
    )
    add_para(
        doc,
        "The scientific and practical value of the thesis lies in providing an audited, cross-platform "
        "evidence base for digital health communication on physical activity. The findings can support "
        "public health guidance, improve health literacy education, and inform platform-level quality "
        "interventions. The thesis is structured as follows. Chapter 1 reviews the literature and "
        "theoretical background. Chapter 2 presents the methodology. Chapter 3 reports the research "
        "findings. Chapter 4 discusses the results, limitations, and implications. The final "
        "unnumbered sections present the conclusions, recommendations, references, and annexes.",
        first_line_indent_cm=1.25,
    )

    path = os.path.join(CHAPTERS, "introduction.docx")
    doc.save(path)
    print(f"Introduction saved to {path}")


if __name__ == "__main__":
    generate_flyleaf()
    generate_abstract()
    generate_acknowledgments()
    generate_abbreviations()
    generate_introduction()
    print("\nAll front matter generated successfully.")
