#!/usr/bin/env python3
"""
Generate Chapter 3: Research Findings for Master's Thesis
Lithuanian Sports University
Author: Ayokunle Ademola-John
"""

import subprocess
import sys

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# HELPER FUNCTIONS (matching generate_chapter2.py)
# ============================================================

def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_table_borders(table):
    """Apply borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def format_run(run, bold=False, italic=False, size=12, font_name="Times New Roman", color=None):
    """Format a run with specified properties."""
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_formatted_paragraph(doc, text, bold=False, italic=False, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=0, space_after=6, line_spacing=1.5,
                             first_line_indent=None):
    """Add a fully formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    format_run(run, bold=bold, italic=italic, size=size)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_heading_custom(doc, text, level=1, size=14, space_before=12, space_after=6):
    """Add a custom heading with Times New Roman formatting."""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    format_run(run, bold=True, size=size)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=1.5)
    return p


def add_apa_table_title(doc, table_number, title_text):
    """Add APA-style table title (number in bold, title in bold italic)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run(f"Table {table_number}")
    format_run(run1, bold=True, italic=False, size=12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(title_text)
    format_run(run2, bold=False, italic=True, size=12)
    add_paragraph_spacing(p, before=12, after=0, line_spacing=1.5)
    add_paragraph_spacing(p2, before=0, after=6, line_spacing=1.5)
    return p, p2


def format_table_cell(cell, text, bold=False, italic=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Format a table cell with text."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    format_run(run, bold=bold, italic=italic, size=size)
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15


def add_mixed_paragraph(doc, parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=0, space_after=6, line_spacing=1.5,
                         first_line_indent=None):
    """
    Add a paragraph with mixed formatting.
    parts: list of tuples (text, bold, italic)
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    for text, bold, italic in parts:
        run = p.add_run(text)
        format_run(run, bold=bold, italic=italic, size=12)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_apa_table_note(doc, note_text):
    """Add an APA-style table note."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_label = p.add_run("Note. ")
    format_run(run_label, bold=False, italic=True, size=10)
    run_text = p.add_run(note_text)
    format_run(run_text, bold=False, italic=False, size=10)
    add_paragraph_spacing(p, before=0, after=12, line_spacing=1.5)
    return p


def create_apa_table(doc, headers, rows, col_widths=None):
    """Create an APA-style table with top/bottom borders and header underline."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for j, header in enumerate(headers):
        format_table_cell(table.rows[0].cells[j], header, bold=True, size=10,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            format_table_cell(table.rows[i + 1].cells[j], str(cell_text), size=10,
                              alignment=align)

    # Apply APA-style borders: top and bottom of table, plus header underline
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Add bottom border to header row cells
    for cell in table.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

    # Set column widths if provided
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, width in enumerate(col_widths):
                row.cells[j].width = Inches(width)

    return table


# ============================================================
# MAIN DOCUMENT GENERATION
# ============================================================

def generate_chapter3():
    doc = Document()

    # ---- Page Setup ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ================================================================
    # CHAPTER HEADING
    # ================================================================
    add_heading_custom(doc, "3. RESEARCH FINDINGS", level=1, size=14, space_before=0, space_after=12)

    add_formatted_paragraph(
        doc,
        "This chapter presents the results of the cross-sectional content analysis of 200 "
        "physical activity information sources sampled across four platforms: traditional websites "
        "(via Google search), YouTube, TikTok, and Instagram. Results are organized by research "
        "question, beginning with descriptive characteristics of the sample, followed by inter-rater "
        "reliability, DISCERN and JAMA benchmark scores across platforms, creator type analysis, "
        "engagement-quality relationships, and content format patterns.",
        first_line_indent=1.25
    )

    # ================================================================
    # 3.1 SAMPLE CHARACTERISTICS
    # ================================================================
    add_heading_custom(doc, "3.1 Sample Characteristics", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "A total of 200 items were sampled, with 50 items from each of the four platforms "
        "(traditional websites, YouTube, TikTok, and Instagram), derived from the top 10 "
        "results for each of the five search terms. All items met the inclusion criteria of "
        "being publicly accessible, English-language, substantive physical activity content "
        "posted or updated within the preceding three years.",
        first_line_indent=1.25
    )

    add_formatted_paragraph(
        doc,
        "Creator type distribution varied substantially across platforms. Traditional websites "
        "were predominantly produced by organizations (39 of 50, 78.0%), reflecting the dominance "
        "of institutional health information providers in Google search results. YouTube content "
        "was most frequently created by fitness influencers (19 of 50, 38.0%), followed by "
        "certified fitness professionals (14 of 50, 28.0%). TikTok similarly featured fitness "
        "influencers as the largest creator category (21 of 50, 42.0%). Instagram content was "
        "predominantly from organizations (29 of 50, 58.0%). Healthcare professionals constituted "
        "a small proportion across all platforms (total n = 14, 7.0%).",
        first_line_indent=1.25
    )

    add_formatted_paragraph(
        doc,
        "Content format was largely determined by platform affordances. Website content consisted "
        "entirely of text articles. YouTube was dominated by long-form video exceeding five minutes "
        "(45 of 50, 90.0%). TikTok featured predominantly short-form video under 60 seconds "
        "(18 of 50, 36.0% of captured format data). All 50 Instagram items were carousel posts "
        "(image and caption combinations). Table 1 summarizes the sample distribution by platform, "
        "creator type, and content format.",
        first_line_indent=1.25
    )

    # Table 1: Sample Characteristics
    add_apa_table_title(doc, 1, "Sample Characteristics by Platform (N = 200)")

    headers_t1 = ["Characteristic", "Website\n(n = 50)", "YouTube\n(n = 50)", "TikTok\n(n = 50)", "Instagram\n(n = 50)", "Total\n(N = 200)"]
    rows_t1 = [
        ["Creator type", "", "", "", "", ""],
        ["  Organization", "39 (78.0)", "5 (10.0)", "12 (24.0)", "29 (58.0)", "85 (42.5)"],
        ["  Certified fitness prof.", "5 (10.0)", "14 (28.0)", "8 (16.0)", "7 (14.0)", "34 (17.0)"],
        ["  Healthcare professional", "4 (8.0)", "5 (10.0)", "3 (6.0)", "2 (4.0)", "14 (7.0)"],
        ["  Fitness influencer", "1 (2.0)", "19 (38.0)", "21 (42.0)", "2 (4.0)", "43 (21.5)"],
        ["  General user", "1 (2.0)", "7 (14.0)", "6 (12.0)", "10 (20.0)", "24 (12.0)"],
        ["Content format", "", "", "", "", ""],
        ["  Text article", "50 (100)", "\u2014", "\u2014", "\u2014", "50 (25.0)"],
        ["  Long video (>5 min)", "\u2014", "45 (90.0)", "\u2014", "\u2014", "45 (22.5)"],
        ["  Short video (<60 s)", "\u2014", "5 (10.0)", "18 (36.0)", "\u2014", "23 (11.5)"],
        ["  Carousel", "\u2014", "\u2014", "\u2014", "50 (100)", "50 (25.0)"],
    ]
    create_apa_table(doc, headers_t1, rows_t1, col_widths=[1.8, 1.0, 1.0, 1.0, 1.0, 1.0])
    add_apa_table_note(doc, "Values are n (%). Dashes indicate format not applicable to platform. "
                       "TikTok format data were captured for a subset of items.")

    # ================================================================
    # 3.2 INTER-RATER RELIABILITY
    # ================================================================
    add_heading_custom(doc, "3.2 Inter-Rater Reliability", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "A randomly selected 20% subsample (n = 40, 10 per platform) was independently scored "
        "by a second rater to assess inter-rater reliability. For the DISCERN instrument, the "
        "intraclass correlation coefficient (ICC) for total scores was 0.994 (95% CI: 0.990\u20131.000), "
        "indicating excellent agreement between raters. For the JAMA benchmark criteria, Cohen\u2019s "
        "kappa was 1.000 for all four criteria (authorship, attribution, disclosure, and currency), "
        "reflecting perfect inter-rater agreement. These values exceeded the pre-established "
        "thresholds of ICC \u2265 0.75 and \u03ba \u2265 0.61, confirming the reliability of the scoring "
        "procedures.",
        first_line_indent=1.25
    )

    # ================================================================
    # 3.3 DISCERN SCORES BY PLATFORM (RQ1)
    # ================================================================
    add_heading_custom(doc, "3.3 DISCERN Scores by Platform", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "Research Question 1 asked how physical activity information quality differs between "
        "traditional websites and social media platforms as measured by DISCERN and JAMA benchmarks. "
        "This section reports DISCERN findings; JAMA results follow in Section 3.4.",
        first_line_indent=1.25
    )

    # Overall descriptives
    add_formatted_paragraph(
        doc,
        "Across all 200 items, the overall mean DISCERN total score was 36.30 (SD = 11.42), "
        "with a median of 35.50 (range 16\u201365). This places the aggregate sample at the boundary "
        "between the \u201cpoor\u201d and \u201cfair\u201d quality categories according to established DISCERN "
        "interpretation thresholds.",
        first_line_indent=1.25
    )

    # Platform-specific
    add_formatted_paragraph(
        doc,
        "Substantial variation was observed across platforms. Traditional websites achieved "
        "the highest median DISCERN score (Mdn = 48.00, IQR = 42.25\u201355.00), followed by "
        "YouTube (Mdn = 39.00, IQR = 36.00\u201342.75), TikTok (Mdn = 31.00, IQR = 29.00\u201337.00), "
        "and Instagram (Mdn = 21.00, IQR = 18.25\u201329.75). Table 2 presents the full descriptive "
        "statistics.",
        first_line_indent=1.25
    )

    # Table 2: DISCERN Descriptive Statistics
    add_apa_table_title(doc, 2, "DISCERN Total Scores by Platform")

    headers_t2 = ["Platform", "n", "M", "SD", "Mdn", "IQR", "Range", "Quality category"]
    rows_t2 = [
        ["Website", "50", "48.70", "8.17", "48.00", "42.25\u201355.00", "31\u201365", "Fair\u2013Good"],
        ["YouTube", "50", "40.36", "6.21", "39.00", "36.00\u201342.75", "28\u201363", "Fair"],
        ["TikTok", "50", "32.38", "6.39", "31.00", "29.00\u201337.00", "20\u201351", "Poor"],
        ["Instagram", "50", "23.78", "5.83", "21.00", "18.25\u201329.75", "16\u201339", "Very Poor\u2013Poor"],
        ["Overall", "200", "36.30", "11.42", "35.50", "24.00\u201344.00", "16\u201365", "Poor\u2013Fair"],
    ]
    create_apa_table(doc, headers_t2, rows_t2, col_widths=[0.9, 0.4, 0.55, 0.5, 0.6, 1.1, 0.75, 1.2])
    add_apa_table_note(doc, "Quality categories based on DISCERN thresholds: 16\u201326 = Very Poor, "
                       "27\u201338 = Poor, 39\u201350 = Fair, 51\u201362 = Good, 63\u201380 = Excellent.")

    # Kruskal-Wallis
    add_formatted_paragraph(
        doc,
        "A Kruskal-Wallis H test confirmed that DISCERN total scores differed significantly "
        "across the four platforms, H(3) = 136.58, p < .001, \u03b7\u00b2 = .68, indicating a large "
        "effect. All six pairwise comparisons using Dunn\u2019s post-hoc test with Bonferroni "
        "correction were statistically significant (p < .05). Table 3 presents the pairwise "
        "comparisons with effect sizes.",
        first_line_indent=1.25
    )

    # Table 3: Pairwise Comparisons
    add_apa_table_title(doc, 3, "Dunn\u2019s Pairwise Comparisons of DISCERN Total Scores")

    headers_t3 = ["Comparison", "z", "p (adjusted)", "r"]
    rows_t3 = [
        ["Website vs. Instagram", "\u2014", "< .001", "\u22120.995"],
        ["Website vs. TikTok", "\u2014", "< .001", "\u22120.878"],
        ["Website vs. YouTube", "\u2014", "< .05", "\u22120.584"],
        ["YouTube vs. Instagram", "\u2014", "< .001", "\u22120.974"],
        ["YouTube vs. TikTok", "\u2014", "< .001", "\u22120.624"],
        ["TikTok vs. Instagram", "\u2014", "< .001", "\u22120.641"],
    ]
    create_apa_table(doc, headers_t3, rows_t3, col_widths=[2.2, 0.8, 1.2, 1.0])
    add_apa_table_note(doc, "All comparisons significant at p < .05 after Bonferroni correction. "
                       "r = rank-biserial correlation effect size. All effect sizes are large (|r| > .50).")

    # DISCERN quality categories
    add_formatted_paragraph(
        doc,
        "When DISCERN scores were categorized, a clear quality gradient emerged. Among website "
        "items, 44.0% were rated as Good and 46.0% as Fair. YouTube content was predominantly "
        "Fair (48.0%) or Poor (46.0%). TikTok items were largely Poor (70.0%), with 14.0% rated "
        "Very Poor. Instagram exhibited the lowest quality distribution, with 60.0% of items "
        "rated Very Poor and 40.0% rated Poor (see Figure 1).",
        first_line_indent=1.25
    )

    # DISCERN subscales
    add_formatted_paragraph(
        doc,
        "Analysis of DISCERN subscales revealed parallel patterns. For Section 1 (reliability "
        "of the publication; items Q1\u2013Q8), median scores were: Website = 26.00, YouTube = 22.00, "
        "TikTok = 18.00, and Instagram = 13.00. For Section 2 (quality of treatment choice "
        "information; items Q9\u2013Q15), median scores were: Website = 19.00, YouTube = 14.00, "
        "TikTok = 12.00, and Instagram = 7.00. Both subscales thus exhibited the same rank "
        "ordering of platforms as the total score.",
        first_line_indent=1.25
    )

    # ================================================================
    # 3.4 JAMA COMPLIANCE BY PLATFORM (RQ1)
    # ================================================================
    add_heading_custom(doc, "3.4 JAMA Benchmark Compliance by Platform", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "Overall JAMA benchmark compliance was low across all platforms. The mean JAMA total "
        "score (range 0\u20134) did not differ significantly across platforms, H(3) = 5.43, "
        "p = .143. Mean scores were: Website = 1.44, YouTube = 1.54, TikTok = 1.44, and "
        "Instagram = 1.20. However, analysis of individual JAMA criteria revealed significant "
        "platform differences for three of the four benchmarks. Table 4 presents the compliance "
        "rates for each criterion by platform.",
        first_line_indent=1.25
    )

    # Table 4: JAMA Compliance
    add_apa_table_title(doc, 4, "JAMA Benchmark Compliance by Platform")

    headers_t4 = ["JAMA Criterion", "Website\n% (n)", "YouTube\n% (n)", "TikTok\n% (n)", "Instagram\n% (n)", "\u03c7\u00b2", "p"]
    rows_t4 = [
        ["Authorship", "30 (15)", "38 (19)", "30 (15)", "16 (8)", "6.16", ".104"],
        ["Attribution", "30 (15)", "16 (8)", "0 (0)", "4 (2)", "25.01", "< .001"],
        ["Disclosure", "16 (8)", "0 (0)", "14 (7)", "0 (0)", "16.36", ".001"],
        ["Currency", "68 (34)", "100 (50)", "100 (50)", "100 (50)", "52.17", "< .001"],
    ]
    create_apa_table(doc, headers_t4, rows_t4, col_widths=[1.2, 0.9, 0.9, 0.9, 1.0, 0.6, 0.7])
    add_apa_table_note(doc, "Values are % (n). Chi-square tests with df = 3. "
                       "Bold p values indicate statistical significance at p < .05.")

    add_formatted_paragraph(
        doc,
        "Authorship identification did not differ significantly across platforms, \u03c7\u00b2(3) = 6.16, "
        "p = .104, with compliance rates ranging from 16.0% (Instagram) to 38.0% (YouTube). "
        "Attribution of sources differed significantly, \u03c7\u00b2(3) = 25.01, p < .001, with websites "
        "demonstrating the highest rate (30.0%) and TikTok the lowest (0.0%). Disclosure of "
        "conflicts of interest also differed significantly, \u03c7\u00b2(3) = 16.36, p = .001, with no "
        "YouTube or Instagram items meeting this criterion. Currency compliance differed "
        "significantly, \u03c7\u00b2(3) = 52.17, p < .001, driven by the lower rate for websites (68.0%) "
        "relative to all three social media platforms (100.0% each), where post dates are "
        "automatically displayed.",
        first_line_indent=1.25
    )

    # ================================================================
    # 3.5 CREATOR TYPE ANALYSIS (RQ2)
    # ================================================================
    add_heading_custom(doc, "3.5 Creator Type Analysis", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "Research Question 2 examined whether creator type moderates information quality scores. "
        "Creators were classified as professionals (healthcare professionals and certified fitness "
        "professionals, n = 48) or non-professionals (fitness influencers, general users, and "
        "organizations, n = 152).",
        first_line_indent=1.25
    )

    add_formatted_paragraph(
        doc,
        "Professional creators produced content with significantly higher DISCERN scores "
        "(Mdn = 39.50) compared to non-professional creators (Mdn = 35.00), U = 4,392, "
        "p = .033, r = \u22120.204, representing a small effect size. When examined by individual "
        "creator type, a Kruskal-Wallis test confirmed significant differences, H(4) = 19.25, "
        "p < .001. Certified fitness professionals achieved the highest median score (Mdn = 41.00), "
        "followed by organizations (Mdn = 40.00), healthcare professionals (Mdn = 37.50), "
        "fitness influencers (Mdn = 35.00), and general users (Mdn = 28.50). Table 5 presents "
        "the full breakdown.",
        first_line_indent=1.25
    )

    # Table 5: Creator Type
    add_apa_table_title(doc, 5, "DISCERN Total Scores and JAMA Compliance by Creator Type")

    headers_t5 = ["Creator type", "n", "DISCERN Mdn", "DISCERN IQR", "JAMA Mdn"]
    rows_t5 = [
        ["Professional", "48", "39.50", "\u2014", "2.00"],
        ["  Healthcare professional", "14", "37.50", "\u2014", "\u2014"],
        ["  Certified fitness prof.", "34", "41.00", "\u2014", "\u2014"],
        ["Non-professional", "152", "35.00", "\u2014", "1.00"],
        ["  Organization", "85", "40.00", "\u2014", "\u2014"],
        ["  Fitness influencer", "43", "35.00", "\u2014", "\u2014"],
        ["  General user", "24", "28.50", "\u2014", "\u2014"],
    ]
    create_apa_table(doc, headers_t5, rows_t5, col_widths=[1.8, 0.5, 1.1, 1.1, 0.9])
    add_apa_table_note(doc, "Professional vs. non-professional DISCERN comparison: U = 4,392, p = .033, "
                       "r = \u22120.204. JAMA comparison: p < .001.")

    add_formatted_paragraph(
        doc,
        "Professional creators also demonstrated significantly higher JAMA benchmark compliance "
        "(Mdn = 2.00) compared to non-professional creators (Mdn = 1.00), p < .001. This "
        "pattern was consistent across platforms, though the small number of professional "
        "creators on some platforms limits subgroup analysis.",
        first_line_indent=1.25
    )

    # ================================================================
    # 3.6 ENGAGEMENT-QUALITY RELATIONSHIP (RQ3)
    # ================================================================
    add_heading_custom(doc, "3.6 Engagement-Quality Relationship", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "Research Question 3 investigated the relationship between engagement metrics and "
        "information quality. Spearman\u2019s rank correlations were computed between DISCERN total "
        "scores and three engagement indicators (views, likes, and comments) across the social "
        "media platforms. Table 6 presents the correlation matrix.",
        first_line_indent=1.25
    )

    # Table 6: Engagement Correlations
    add_apa_table_title(doc, 6, "Spearman\u2019s Correlations Between Engagement Metrics and DISCERN Scores")

    headers_t6 = ["Metric", "Overall \u03c1", "p", "YouTube \u03c1", "p", "TikTok \u03c1", "p", "Instagram \u03c1", "p"]
    rows_t6 = [
        ["Views", "0.017", ".894", "\u22120.352", ".018", "0.522", ".026", "\u2014", "ns"],
        ["Likes", "0.549", "< .001", "\u22120.338", ".023", "0.475", ".003", "\u2014", "ns"],
        ["Comments", "0.520", "< .001", "\u2014", "\u2014", "\u2014", "\u2014", "\u2014", "ns"],
    ]
    create_apa_table(doc, headers_t6, rows_t6, col_widths=[0.8, 0.65, 0.55, 0.7, 0.55, 0.65, 0.55, 0.75, 0.4])
    add_apa_table_note(doc, "Dashes indicate data not available or not computed. "
                       "ns = not significant. Website engagement data were not collected.")

    add_formatted_paragraph(
        doc,
        "At the aggregate level, the correlation between views and DISCERN scores was negligible "
        "and non-significant (\u03c1 = .017, p = .894). In contrast, both likes (\u03c1 = .549, p < .001) "
        "and comments (\u03c1 = .520, p < .001) demonstrated strong positive correlations with "
        "information quality across platforms.",
        first_line_indent=1.25
    )

    add_formatted_paragraph(
        doc,
        "Platform-specific analysis revealed divergent patterns. On YouTube, views and likes "
        "were negatively correlated with DISCERN scores (\u03c1 = \u2212.352, p = .018 and \u03c1 = \u2212.338, "
        "p = .023, respectively), suggesting that more popular YouTube content tended to be of "
        "lower quality. Conversely, on TikTok, both views (\u03c1 = .522, p = .026) and likes "
        "(\u03c1 = .475, p = .003) were positively associated with quality. No significant "
        "engagement-quality correlations were observed for Instagram.",
        first_line_indent=1.25
    )

    # ================================================================
    # 3.7 CONTENT FORMAT PATTERNS
    # ================================================================
    add_heading_custom(doc, "3.7 Content Format Patterns", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "Content format was associated with information quality. Long-form video content "
        "(>5 minutes) achieved a median DISCERN score of 39.00, compared to 29.00 for short-form "
        "video (<60 seconds) and 21.00 for carousel posts. These differences are largely "
        "confounded with platform, as content format was heavily determined by platform affordances. "
        "However, the pattern suggests that formats allowing for greater depth of information "
        "presentation were associated with higher quality scores.",
        first_line_indent=1.25
    )

    add_formatted_paragraph(
        doc,
        "An analysis of search term effects on DISCERN scores revealed no significant differences, "
        "H(4) = 5.32, p = .256, indicating that the five search terms yielded comparable quality "
        "distributions. This finding supports the robustness of the results across different "
        "physical activity information queries.",
        first_line_indent=1.25
    )

    # ================================================================
    # SUMMARY
    # ================================================================
    add_heading_custom(doc, "3.8 Summary of Key Findings", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "In summary, the principal findings of this analysis are as follows. First, a clear "
        "quality gradient was observed across platforms, with traditional websites scoring highest "
        "on DISCERN and Instagram scoring lowest, and all pairwise differences reaching statistical "
        "significance with large effect sizes (\u03b7\u00b2 = .68). Second, JAMA benchmark compliance was "
        "uniformly low across all platforms, with no significant overall difference, though "
        "individual criteria varied by platform. Third, professional creators produced significantly "
        "higher-quality content than non-professionals on both DISCERN and JAMA measures, albeit "
        "with a small effect size for DISCERN (r = \u22120.204). Fourth, the engagement-quality "
        "relationship was metric- and platform-dependent: views showed no overall association with "
        "quality, while likes and comments showed strong positive correlations; however, YouTube "
        "exhibited negative engagement-quality correlations while TikTok showed positive ones. "
        "These findings are discussed in the context of existing literature in Chapter 4.",
        first_line_indent=1.25
    )

    # ---- Save ----
    output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "chapters")
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "chapter3_findings.docx")
    doc.save(output_path)
    print(f"Chapter 3 saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    generate_chapter3()
