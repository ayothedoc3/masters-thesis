"""
Assemble the complete thesis from all components into a single .docx file.

Front matter sections are intentionally excluded from the automatic table of contents.
The TOC begins with Introduction and ends with Annexes to match LSU regulations.
"""
from __future__ import annotations

import os
import re
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT = os.path.join(ROOT, "final", "thesis_complete.docx")
FINAL_DIR = os.path.join(ROOT, "final")
os.makedirs(FINAL_DIR, exist_ok=True)


def setup_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        hs.font.size = Pt(14 if level == 1 else 12)
        hs.paragraph_format.line_spacing = 1.5
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc


def add_centered_line(doc: Document, text: str, *, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold


def add_title_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph("")

    add_centered_line(doc, "LITHUANIAN SPORTS UNIVERSITY", bold=True)
    add_centered_line(doc, "MASTER OF PUBLIC HEALTH PROGRAMME", bold=True)
    add_centered_line(doc, "FACULTY OF SPORT MEDICINE", bold=True)

    for _ in range(5):
        doc.add_paragraph("")

    add_centered_line(
        doc,
        "QUALITY OF PHYSICAL ACTIVITY INFORMATION ON TRADITIONAL WEBSITES "
        "VERSUS SOCIAL MEDIA PLATFORMS:",
        bold=True,
        size=14,
    )
    add_centered_line(
        doc,
        "A CROSS-SECTIONAL CONTENT ANALYSIS USING DISCERN AND JAMA BENCHMARKS",
        bold=True,
        size=14,
    )

    for _ in range(4):
        doc.add_paragraph("")

    title_lines = [
        "Final Master's Thesis",
        "",
        "Student: Ayokunle Ademola-John ____________________",
        "Scientific supervisor: Dr. Antanas Usas ____________________",
        "Scientific adviser: _________________________________________",
        "",
        "Kaunas, 2026",
    ]
    for line in title_lines:
        add_centered_line(doc, line)

    doc.add_page_break()


def append_docx(doc: Document, source_path: str, label: str = "") -> None:
    print(f"  Appending: {label or source_path}")
    src = Document(source_path)

    rel_map: dict[str, str] = {}
    for rel in src.part.rels.values():
        if "image" in rel.reltype:
            new_rel = doc.part.relate_to(rel.target_part, rel.reltype)
            rel_map[rel.rId] = new_rel

    for element in src.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag not in ("p", "tbl"):
            continue
        new_elem = deepcopy(element)
        if rel_map:
            for blip in new_elem.iter(qn("a:blip")):
                embed = blip.get(qn("r:embed"))
                if embed and embed in rel_map:
                    blip.set(qn("r:embed"), rel_map[embed])
        body = doc.element.body
        insert_at = len(body)
        if len(body) and body[-1].tag == qn("w:sectPr"):
            insert_at = len(body) - 1
        body.insert(insert_at, new_elem)


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and update field in Word if page numbers do not appear automatically."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def add_toc_page(doc: Document) -> None:
    add_centered_line(doc, "TABLE OF CONTENTS", bold=True, size=14)
    toc_paragraph = doc.add_paragraph()
    toc_paragraph.paragraph_format.line_spacing = 1.5
    add_toc_field(toc_paragraph)
    doc.add_page_break()


def add_annexes_heading(doc: Document) -> None:
    p = doc.add_paragraph("ANNEXES", style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def normalize_structure(doc: Document) -> None:
    front_matter_titles = {
        "CONFIRMATION OF INDEPENDENT COMPOSITION OF THE THESIS",
        "CONFIRMATION OF LIABILITY FOR THE REGULARITY OF THE ENGLISH LANGUAGE",
        "FINAL MASTER'S THESIS SUPERVISOR'S ASSESSMENT",
        "ABSTRACT",
        "SANTRAUKA",
        "ACKNOWLEDGMENTS",
        "LIST OF ABBREVIATIONS",
        "TABLE OF CONTENTS",
    }
    top_level_titles = {"INTRODUCTION", "CONCLUSIONS", "RECOMMENDATIONS", "REFERENCES", "ANNEXES"}

    in_toc = False
    in_annexes = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        if text == "TABLE OF CONTENTS":
            in_toc = True
            paragraph.style = "Normal"
            continue

        if in_toc:
            if text == "CONFIRMATION OF INDEPENDENT COMPOSITION OF THE THESIS":
                in_toc = False
            else:
                continue

        if text == "CHAPTER 4. DISCUSSION":
            paragraph.text = "CHAPTER 4. CONSIDERATIONS"
            text = paragraph.text.strip()

        if text in front_matter_titles:
            paragraph.style = "Normal"
            continue

        if text == "ANNEXES":
            in_annexes = True
            paragraph.style = "Heading 1"
            continue

        if in_annexes:
            paragraph.style = "Normal"
            continue

        if text in top_level_titles or re.match(r"^CHAPTER\s+\d+\.\s+", text):
            paragraph.style = "Heading 1"
            continue

        if re.match(r"^\d+\.\d+\s+", text):
            paragraph.style = "Heading 2"
            continue


def main() -> None:
    print("=" * 60)
    print("ASSEMBLING COMPLETE THESIS")
    print("=" * 60)

    doc = setup_doc()

    add_title_page(doc)
    append_docx(doc, os.path.join(ROOT, "chapters", "flyleaf.docx"), "Flyleaf")
    doc.add_page_break()
    add_toc_page(doc)
    append_docx(doc, os.path.join(ROOT, "chapters", "abstract.docx"), "Abstract")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "chapters", "acknowledgments.docx"), "Acknowledgments")
    append_docx(doc, os.path.join(ROOT, "chapters", "abbreviations.docx"), "Abbreviations")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "chapters", "introduction.docx"), "Introduction")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "chapters", "chapter1_literature_review.docx"), "Chapter 1")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "chapters", "chapter2_methodology.docx"), "Chapter 2")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "chapters", "chapter3_findings.docx"), "Chapter 3")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "chapters", "chapter4_discussion.docx"), "Chapter 4")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "chapters", "chapter5_conclusions.docx"), "Conclusions and Recommendations")
    doc.add_page_break()
    append_docx(doc, os.path.join(ROOT, "references", "references.docx"), "References")
    doc.add_page_break()
    add_annexes_heading(doc)
    doc.add_page_break()

    appendices = [
        ("appendix_a_discern.docx", "Appendix A: DISCERN"),
        ("appendix_b_jama.docx", "Appendix B: JAMA"),
        ("appendix_c_data_summary.docx", "Appendix C: Data Summary"),
        ("appendix_d_analysis_code.docx", "Appendix D: Analysis Code"),
        ("appendix_e_article.docx", "Appendix E: Article"),
        ("appendix_f_conference_evidence.docx", "Appendix F: Conference Evidence"),
        ("appendix_g_consultations_timesheet.docx", "Appendix G: Consultations Timesheet"),
        ("appendix_h_ethics_evidence.docx", "Appendix H: Ethics Evidence"),
    ]
    for idx, (filename, label) in enumerate(appendices, start=1):
        append_docx(doc, os.path.join(ROOT, "appendices", filename), label)
        if idx < len(appendices):
            doc.add_page_break()

    doc.save(OUTPUT)
    normalized = Document(OUTPUT)
    normalize_structure(normalized)
    normalized.save(OUTPUT)

    print(f"\n{'=' * 60}")
    print(f"COMPLETE THESIS SAVED: {OUTPUT}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
