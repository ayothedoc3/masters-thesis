"""
Generate Appendix A: DISCERN Instrument
Appendix B: JAMA Benchmarks Operationalized
Appendix C: Data Summary Tables
Appendix D: Analysis Code Reference
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE)
APPENDIX_DIR = os.path.join(ROOT, "appendices")
os.makedirs(APPENDIX_DIR, exist_ok=True)


def set_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.line_spacing = 1.5
    return p


def make_table(doc, data):
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
            for p in table.rows[i].cells[j].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
                    if i == 0:
                        r.bold = True
    return table


# ============================================================
# APPENDIX A: DISCERN INSTRUMENT
# ============================================================
def generate_appendix_a():
    doc = Document()
    set_style(doc)

    add_heading(doc, "APPENDIX A: DISCERN INSTRUMENT", level=1)
    add_para(doc, "")
    add_para(
        doc,
        "The DISCERN instrument (Charnock et al., 1999) is a validated tool for assessing "
        "the quality of written consumer health information. It consists of 16 questions across "
        "three sections, each scored on a 5-point Likert scale (1 = Definitely No to 5 = Definitely Yes). "
        "The instrument was originally designed for patient information leaflets but has been widely "
        "adapted for online health content assessment.",
    )
    add_para(doc, "")

    add_heading(doc, "Section 1: Reliability of the Publication (Questions 1\u20138)", level=2)
    questions_s1 = [
        ("Q1", "Are the aims clear?", "Does the publication clearly state what it is about and what topics will be covered?"),
        ("Q2", "Does it achieve its aims?", "Does the content deliver on the stated purpose?"),
        ("Q3", "Is it relevant?", "Is the information relevant to the target audience?"),
        ("Q4", "Is it clear what sources of information were used?", "Are references, sources, or evidence cited?"),
        ("Q5", "Is it clear when the information was produced?", "Is the date of creation or last update stated?"),
        ("Q6", "Is it balanced and unbiased?", "Does the content present information objectively, or is it one-sided?"),
        ("Q7", "Does it provide details of additional sources of support?", "Are links to further reading, helplines, or professional resources provided?"),
        ("Q8", "Does it refer to areas of uncertainty?", "Does the content acknowledge limitations, uncertainties, or areas where evidence is lacking?"),
    ]

    for qnum, title, desc in questions_s1:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"{qnum}: {title}")
        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        p2 = doc.add_paragraph()
        p2.paragraph_format.line_spacing = 1.5
        p2.paragraph_format.left_indent = Inches(0.5)
        r2 = p2.add_run(desc)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
        r2.italic = True

    add_para(doc, "")
    add_heading(doc, "Section 2: Quality of Information on Treatment Choices (Questions 9\u201315)", level=2)
    questions_s2 = [
        ("Q9", "Does it describe how each treatment works?", "Does the content explain the mechanism or rationale behind recommended exercises or activities?"),
        ("Q10", "Does it describe the benefits of each treatment?", "Are the expected benefits of the recommended physical activity clearly stated?"),
        ("Q11", "Does it describe the risks of each treatment?", "Are potential risks, contraindications, or side effects of the recommended exercises discussed?"),
        ("Q12", "Does it describe what would happen if no treatment is used?", "Does the content explain the consequences of physical inactivity?"),
        ("Q13", "Does it describe how treatment choices affect overall quality of life?", "Does the content address how physical activity affects wellbeing beyond the specific condition?"),
        ("Q14", "Is it clear that there may be more than one possible treatment choice?", "Does the content present multiple exercise options rather than a single approach?"),
        ("Q15", "Does it provide support for shared decision-making?", "Does the content encourage consulting a professional or making personalised decisions?"),
    ]

    for qnum, title, desc in questions_s2:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        r1 = p.add_run(f"{qnum}: {title}")
        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        p2 = doc.add_paragraph()
        p2.paragraph_format.line_spacing = 1.5
        p2.paragraph_format.left_indent = Inches(0.5)
        r2 = p2.add_run(desc)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
        r2.italic = True

    add_para(doc, "")
    add_heading(doc, "Section 3: Overall Quality Rating (Question 16)", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r1 = p.add_run(
        "Q16: Based on the answers to all of the above questions, rate the overall quality "
        "of the publication as a source of information about treatment choices."
    )
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(12)

    add_para(doc, "")
    add_heading(doc, "Scoring and Interpretation", level=2)
    add_para(
        doc,
        "Each question is scored from 1 (Definitely No) to 5 (Definitely Yes). The total "
        "score ranges from 16 to 80. The following categories are used for interpretation:",
    )
    add_para(doc, "")

    make_table(doc, [
        ["Score Range", "Category", "Interpretation"],
        ["16\u201326", "Very Poor", "Serious or extensive shortcomings"],
        ["27\u201338", "Poor", "Potentially important but not serious shortcomings"],
        ["39\u201350", "Fair", "Some shortcomings but not serious"],
        ["51\u201362", "Good", "Minimal shortcomings"],
        ["63\u201380", "Excellent", "Minimal shortcomings with additional features"],
    ])

    path = os.path.join(APPENDIX_DIR, "appendix_a_discern.docx")
    doc.save(path)
    print(f"Appendix A saved to {path}")


# ============================================================
# APPENDIX B: JAMA BENCHMARKS
# ============================================================
def generate_appendix_b():
    doc = Document()
    set_style(doc)

    add_heading(doc, "APPENDIX B: JAMA BENCHMARKS \u2014 OPERATIONALISATION FOR SOCIAL MEDIA", level=1)
    add_para(doc, "")
    add_para(
        doc,
        "The JAMA benchmarks (Silberg et al., 1997) were originally proposed as minimum standards "
        "for health information published on the World Wide Web. They comprise four accountability "
        "criteria, each assessed as present (Y) or absent (N). This appendix details how each "
        "criterion was operationalised for the four platform types assessed in this study.",
    )
    add_para(doc, "")

    criteria = [
        ("1. Authorship",
         "Authors and contributors should be identified, along with their affiliations and relevant credentials.",
         [("Traditional websites", "Author name and credentials (e.g., MD, PhD, CSCS) stated on the page, in a byline, or in an About the Author section."),
          ("YouTube", "Creator credentials stated in the video description, channel About page, or verbally in the video introduction."),
          ("TikTok", "Creator bio states professional credentials (e.g., Certified Personal Trainer, Physiotherapist) OR account is verified as a professional."),
          ("Instagram", "Creator bio states professional credentials OR profile is a verified professional/business account with stated qualifications.")]),
        ("2. Attribution",
         "References and sources for all content should be listed clearly.",
         [("Traditional websites", "In-text citations, a reference list, bibliography, or hyperlinks to source materials are provided."),
          ("YouTube", "Sources cited in the video description, shown on screen, or verbally referenced with sufficient detail to locate."),
          ("TikTok", "Sources cited in the caption, shown as on-screen text, or referenced verbally (e.g., According to the WHO guidelines)."),
          ("Instagram", "Sources cited in the post caption, shown on carousel slides, or included as tagged references.")]),
        ("3. Disclosure",
         "Website ownership, sponsorship, advertising, underwriting, commercial funding, or potential conflicts of interest should be disclosed.",
         [("Traditional websites", "A disclosure statement, privacy policy identifying sponsorships, or declaration of conflicts of interest is present."),
          ("YouTube", "Video includes a paid promotion disclosure, sponsorship mention, or the creator states potential conflicts of interest."),
          ("TikTok", "Post includes #ad, #sponsored, Paid partnership label, or the creator verbally discloses commercial relationships."),
          ("Instagram", "Post includes Paid partnership with label, #ad, #sponsored tags, or disclosure in the caption.")]),
        ("4. Currency",
         "Dates of content posting and any updates should be clearly displayed.",
         [("Traditional websites", "Publication date or last updated date is visible on the page."),
          ("YouTube", "Upload date is displayed by default on the platform (always Y for YouTube)."),
          ("TikTok", "Post date is displayed by default on the platform (always Y for TikTok)."),
          ("Instagram", "Post date is displayed by default on the platform (always Y for Instagram).")]),
    ]

    for title, definition, platforms in criteria:
        add_heading(doc, title, level=2)
        add_para(doc, f"Definition: {definition}", italic=True)
        add_para(doc, "")
        data = [["Platform", "Operationalisation"]]
        for plat, desc in platforms:
            data.append([plat, desc])
        make_table(doc, data)
        add_para(doc, "")

    add_heading(doc, "Scoring", level=2)
    add_para(
        doc,
        "Each criterion is scored as Y (Yes, criterion met) or N (No, criterion not met). "
        "The JAMA Total score is the count of criteria met, ranging from 0 to 4.",
    )

    path = os.path.join(APPENDIX_DIR, "appendix_b_jama.docx")
    doc.save(path)
    print(f"Appendix B saved to {path}")


# ============================================================
# APPENDIX C: DATA SUMMARY TABLES
# ============================================================
def generate_appendix_c():
    doc = Document()
    set_style(doc)

    add_heading(doc, "APPENDIX C: COMPLETE DATA SUMMARY TABLES", level=1)
    add_para(doc, "")

    add_heading(doc, "C.1 Sample Characteristics", level=2)

    add_para(doc, "Table C1. Creator Type Distribution by Platform", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    make_table(doc, [
        ["Platform", "Healthcare Prof.", "Certified Fitness Prof.", "Fitness Influencer", "General User", "Organisation"],
        ["Website", "0", "7", "0", "4", "39"],
        ["YouTube", "3", "12", "19", "4", "12"],
        ["TikTok", "8", "7", "21", "9", "5"],
        ["Instagram", "3", "8", "3", "7", "29"],
        ["Total", "14", "34", "43", "24", "85"],
    ])
    add_para(doc, "")

    add_para(doc, "Table C2. DISCERN Total Score Descriptive Statistics by Platform", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    make_table(doc, [
        ["Platform", "n", "Mean", "SD", "Median", "Q1", "Q3", "Min", "Max", "IQR"],
        ["Website", "50", "48.70", "8.17", "48.0", "42.25", "55.00", "33", "65", "12.75"],
        ["YouTube", "50", "40.36", "6.21", "39.0", "36.00", "42.75", "31", "57", "6.75"],
        ["TikTok", "50", "32.38", "6.39", "31.0", "29.00", "37.00", "17", "45", "8.00"],
        ["Instagram", "50", "23.78", "5.83", "21.0", "18.25", "29.75", "16", "36", "11.50"],
        ["Overall", "200", "36.30", "11.42", "35.5", "\u2014", "\u2014", "16", "65", "\u2014"],
    ])
    add_para(doc, "")

    add_para(doc, "Table C3. DISCERN Quality Category Distribution by Platform (n and %)", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    make_table(doc, [
        ["Platform", "Very Poor n (%)", "Poor n (%)", "Fair n (%)", "Good n (%)", "Excellent n (%)"],
        ["Website", "0 (0%)", "4 (8%)", "23 (46%)", "22 (44%)", "1 (2%)"],
        ["YouTube", "0 (0%)", "23 (46%)", "24 (48%)", "3 (6%)", "0 (0%)"],
        ["TikTok", "7 (14%)", "35 (70%)", "8 (16%)", "0 (0%)", "0 (0%)"],
        ["Instagram", "30 (60%)", "20 (40%)", "0 (0%)", "0 (0%)", "0 (0%)"],
    ])
    add_para(doc, "")

    add_para(doc, "Table C4. JAMA Benchmark Compliance by Platform (% Yes)", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    make_table(doc, [
        ["Platform", "Authorship", "Attribution", "Disclosure", "Currency", "Mean Total"],
        ["Website", "30%", "30%", "16%", "68%", "1.44"],
        ["YouTube", "38%", "16%", "0%", "100%", "1.54"],
        ["TikTok", "30%", "0%", "14%", "100%", "1.44"],
        ["Instagram", "16%", "4%", "0%", "100%", "1.20"],
    ])
    add_para(doc, "")

    add_heading(doc, "C.2 Pairwise Comparisons", level=2)
    add_para(doc, "Table C5. Pairwise Platform Comparisons for DISCERN Total (Mann-Whitney U with Bonferroni Correction)", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    make_table(doc, [
        ["Comparison", "U", "p (adjusted)", "Effect size (r)", "Sig."],
        ["Website vs YouTube", "1980", "<.0001", "\u22120.584", "***"],
        ["Website vs TikTok", "2348", "<.0001", "\u22120.878", "***"],
        ["Website vs Instagram", "2494", "<.0001", "\u22120.995", "***"],
        ["YouTube vs TikTok", "2030", "<.0001", "\u22120.624", "***"],
        ["YouTube vs Instagram", "2468", "<.0001", "\u22120.974", "***"],
        ["TikTok vs Instagram", "2051", "<.0001", "\u22120.641", "***"],
    ])
    add_para(doc, "Note. *** p < .001. Effect size r computed as Z / sqrt(N).", italic=True)
    add_para(doc, "")

    add_heading(doc, "C.3 Engagement\u2013Quality Correlations", level=2)
    add_para(doc, "Table C6. Spearman Correlations Between Engagement Metrics and DISCERN Total by Platform", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    make_table(doc, [
        ["Platform", "Metric", "n", "rho", "p", "Sig."],
        ["Overall", "Views", "63", ".017", ".894", ""],
        ["Overall", "Likes", "121", ".549", "<.001", "***"],
        ["Overall", "Comments", "111", ".520", "<.001", "***"],
        ["YouTube", "Views", "\u2014", "\u2212.352", ".018", "*"],
        ["YouTube", "Likes", "\u2014", "\u2212.338", ".023", "*"],
        ["YouTube", "Comments", "\u2014", "\u2212.239", ".127", ""],
        ["TikTok", "Views", "\u2014", ".522", ".026", "*"],
        ["TikTok", "Likes", "\u2014", ".475", ".003", "**"],
        ["TikTok", "Comments", "\u2014", ".262", ".094", ""],
        ["Instagram", "Likes", "\u2014", "\u2212.038", ".817", ""],
        ["Instagram", "Comments", "\u2014", ".065", ".747", ""],
    ])
    add_para(doc, "Note. * p < .05, ** p < .01, *** p < .001.", italic=True)

    path = os.path.join(APPENDIX_DIR, "appendix_c_data_summary.docx")
    doc.save(path)
    print(f"Appendix C saved to {path}")


# ============================================================
# APPENDIX D: ANALYSIS CODE
# ============================================================
def generate_appendix_d():
    doc = Document()
    set_style(doc)

    add_heading(doc, "APPENDIX D: STATISTICAL ANALYSIS CODE", level=1)
    add_para(doc, "")
    add_para(
        doc,
        "The complete statistical analysis was conducted using Python 3.13 with the following "
        "packages: pandas, scipy, pingouin, scikit-posthocs, matplotlib, and seaborn. The full "
        "analysis script (run_analysis.py) is available in the project repository. Key analysis "
        "functions are summarised below.",
    )
    add_para(doc, "")

    add_heading(doc, "D.1 Software Environment", level=2)
    env_items = [
        "Python 3.13",
        "pandas 2.2.x \u2014 data manipulation",
        "scipy 1.14.x \u2014 Kruskal-Wallis H, Mann-Whitney U, Chi-square, Spearman correlation",
        "pingouin 0.5.x \u2014 Intraclass Correlation Coefficient (ICC)",
        "scikit-posthocs 0.9.x \u2014 Dunn\u2019s post-hoc test with Bonferroni correction",
        "matplotlib 3.9.x + seaborn 0.13.x \u2014 figure generation",
    ]
    for item in env_items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    add_para(doc, "")
    add_heading(doc, "D.2 Key Statistical Tests", level=2)

    tests = [
        ("Inter-rater reliability",
         "ICC(3,1) computed via pingouin.intraclass_corr() for DISCERN total scores. "
         "Cohen\u2019s kappa calculated manually for binary JAMA criteria using observed vs. expected agreement."),
        ("Platform comparisons (DISCERN)",
         "Kruskal-Wallis H test (scipy.stats.kruskal) followed by Mann\u2013Whitney U pairwise "
         "comparisons with Bonferroni correction. Effect size: \u03b7\u00b2 = (H \u2212 k + 1) / (N \u2212 k), "
         "where H is the Kruskal-Wallis statistic, k the number of groups, and N the total "
         "sample size (Tomczak & Tomczak, 2014). Pairwise effect sizes reported as rank-biserial "
         "correlation r."),
        ("Platform comparisons (JAMA)",
         "Chi-square test of independence (scipy.stats.chi2_contingency) for each JAMA criterion. "
         "Fisher\u2019s exact test used when expected cell counts < 5. Effect size reported as "
         "Cram\u00e9r\u2019s V."),
        ("Creator type analysis",
         "Mann-Whitney U test (scipy.stats.mannwhitneyu) comparing professional vs. non-professional "
         "creators. Effect size: r = Z / sqrt(N)."),
        ("Engagement-quality correlations",
         "Spearman\u2019s rank correlation (scipy.stats.spearmanr) between engagement metrics and "
         "DISCERN total, both overall and stratified by platform."),
    ]

    for title, desc in tests:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{title}. ")
        r1.bold = True
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(12)
        r2 = p.add_run(desc)
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.5

    add_para(doc, "")
    add_heading(doc, "D.3 Reproducibility", level=2)
    add_para(
        doc,
        "All data, scripts, and analysis outputs are maintained in a version-controlled Git "
        "repository available at https://github.com/ayothedoc3/masters-thesis "
        "(accessed 23 March 2026). The repository snapshot current at the time of submission "
        "contains the materials supporting the submitted thesis. The complete analysis can be reproduced by "
        "running: python analysis/run_analysis.py from the project root directory. The final audited "
        "input files are data/csv/master_dataset_corrected.csv and data/csv/second_rater_scores.csv. "
        "The script normalises common content-format label variants before analysis and writes all "
        "output files (figures, tables, and analysis_report.txt) to analysis/output/.",
    )

    path = os.path.join(APPENDIX_DIR, "appendix_d_analysis_code.docx")
    doc.save(path)
    print(f"Appendix D saved to {path}")


if __name__ == "__main__":
    generate_appendix_a()
    generate_appendix_b()
    generate_appendix_c()
    generate_appendix_d()
    print("\nAll 4 appendices created successfully.")
