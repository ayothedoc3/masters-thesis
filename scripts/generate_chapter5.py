#!/usr/bin/env python3
"""
Generate the final unnumbered thesis sections:
- Conclusions
- Recommendations

These sections are kept concise to match LSU regulations.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_PATH = os.path.join(ROOT, "chapters", "chapter5_conclusions.docx")


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)


def add_numbered_paragraph(doc: Document, number: int, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0)
    run_num = p.add_run(f"{number}. ")
    run_num.font.name = "Times New Roman"
    run_num.font.size = Pt(12)
    run_num.bold = True
    run_txt = p.add_run(text)
    run_txt.font.name = "Times New Roman"
    run_txt.font.size = Pt(12)


def build_document() -> str:
    doc = Document()
    set_style(doc)

    add_heading(doc, "CONCLUSIONS")
    conclusions = [
        (
            "The quality of physical activity information differed significantly across platform "
            "types. Traditional websites provided the highest-quality information, followed by "
            "YouTube, TikTok, and Instagram."
        ),
        (
            "Transparency and accountability were limited across all platforms. JAMA benchmark "
            "compliance was generally low, indicating frequent omission of authorship, attribution, "
            "disclosure, and currency information."
        ),
        (
            "Creator type was associated with information quality. Content produced by healthcare "
            "professionals and certified fitness professionals scored higher than content produced "
            "by general users and many influencer-led accounts."
        ),
        (
            "Engagement indicators did not function as reliable markers of quality. Within YouTube, "
            "higher view counts and like counts were associated with lower DISCERN scores."
        ),
        (
            "The findings support the need for stronger evidence-based communication on social media "
            "and for more explicit quality safeguards in digital physical activity information environments."
        ),
    ]
    for idx, text in enumerate(conclusions, start=1):
        add_numbered_paragraph(doc, idx, text)

    add_heading(doc, "RECOMMENDATIONS")
    recommendations = [
        (
            "Public health institutions should continue to direct users to evidence-based websites "
            "as the primary reference point for physical activity guidance."
        ),
        (
            "Healthcare professionals and certified fitness professionals should strengthen their "
            "presence on short-form social media and consistently provide credentials, sources, "
            "risk information, and disclosure statements."
        ),
        (
            "Platform providers should introduce structured credibility signals for health content, "
            "including visible source-citation and sponsorship-disclosure fields."
        ),
        (
            "Health literacy education should explicitly teach users that views, likes, and other "
            "engagement metrics do not guarantee information quality."
        ),
        (
            "Future institutional and platform initiatives should prioritize validation of quality "
            "assessment tools designed specifically for multimedia and short-form health content."
        ),
    ]
    for idx, text in enumerate(recommendations, start=1):
        add_numbered_paragraph(doc, idx, text)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Chapter 5 saved to: {OUTPUT_PATH}")
    return OUTPUT_PATH


if __name__ == "__main__":
    build_document()
