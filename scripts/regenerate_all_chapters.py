#!/usr/bin/env python3
"""
regenerate_all_chapters.py
Generates Chapters 3, 4, 5 and Appendix C for Ayokunle Ademola-John's MPH thesis.
All statistics are hardcoded from the final analysis output.
LSU formatting: Times New Roman 12pt, A4, 1" margins, 1.5 spacing, justified.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
BASE = r"c:\Users\ayoth\Downloads\Masters Thesis"
CHAPTERS = os.path.join(BASE, "chapters")
APPENDICES = os.path.join(BASE, "appendices")
FIGURES = os.path.join(BASE, "analysis", "output", "figures")

FIG_PATHS = {
    1: os.path.join(FIGURES, "fig1_discern_boxplot.png"),
    2: os.path.join(FIGURES, "fig2_discern_categories.png"),
    3: os.path.join(FIGURES, "fig3_jama_heatmap.png"),
    4: os.path.join(FIGURES, "fig4_creator_type_discern.png"),
    5: os.path.join(FIGURES, "fig5_engagement_vs_quality.png"),
    6: os.path.join(FIGURES, "fig6_discern_questions_heatmap.png"),
}

# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def set_a4_margins(doc):
    """Set A4 page size with 1-inch margins."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def set_run_font(run, size=12, bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Ensure Times New Roman for East Asian fallback
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f"<w:rPr {nsdecls('w')}/>")
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")


def set_paragraph_spacing(paragraph, spacing=1.5, space_after=Pt(6)):
    """Set line spacing and paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.line_spacing = spacing
    pf.space_after = space_after
    pf.space_before = Pt(0)


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, spacing=1.5, space_after=Pt(12))
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, spacing=1.5, space_after=Pt(10))
    run = p.add_run(text)
    set_run_font(run, size=13, bold=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, size=12)
    return p


def add_body_italic(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_spacing(p, spacing=1.5)
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)
    return p


def add_figure(doc, fig_num, caption_text):
    """Insert a figure image and its caption."""
    path = FIG_PATHS.get(fig_num)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
        set_paragraph_spacing(p, spacing=1.0, space_after=Pt(4))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, spacing=1.0, space_after=Pt(12))
    run = cap.add_run(caption_text)
    set_run_font(run, size=11, italic=True)


def format_table(table):
    """Apply consistent formatting to a table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(paragraph, spacing=1.0, space_after=Pt(2))
                for run in paragraph.runs:
                    set_run_font(run, size=10)


def shade_header_row(table):
    """Shade the first row of a table light grey."""
    for cell in table.rows[0].cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=10, bold=True)


def add_table_caption(doc, text):
    """Add a table caption in bold italic."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, spacing=1.5, space_after=Pt(4))
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, italic=True)


def create_doc():
    doc = Document()
    set_a4_margins(doc)
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


# ============================================================================
#  CHAPTER 3 — RESEARCH FINDINGS
# ============================================================================

def build_chapter3():
    doc = create_doc()

    add_heading1(doc, "CHAPTER 3. RESEARCH FINDINGS")

    # ---- 3.1 Sample Characteristics ----
    add_heading2(doc, "3.1 Sample Characteristics")

    add_body(doc,
        "A total of 200 items were sampled across four platforms: traditional websites "
        "(n = 50), YouTube (n = 50), TikTok (n = 50), and Instagram (n = 50). Each platform "
        "contributed 10 results from each of the five search terms, yielding a balanced "
        "representation across content domains. All 200 items met the predefined inclusion "
        "criteria of being publicly accessible, English-language, substantive physical activity "
        "content published or updated within the past three years."
    )

    add_body(doc,
        "Creator types were classified into five categories. As shown in Table 1, "
        "organisations constituted the largest group (n = 85, 42.5%), followed by certified "
        "fitness professionals (n = 31, 15.5%), general users (n = 30, 15.0%), fitness "
        "influencers (n = 27, 13.5%), and healthcare professionals (n = 17, 8.5%). Notably, "
        "organisations were particularly prevalent among traditional websites (n = 34, 68.0%) "
        "and YouTube (n = 25, 50.0%), reflecting the dominance of institutional content on "
        "these platforms. In contrast, TikTok and Instagram were characterised by greater "
        "diversity in creator types, with higher proportions of fitness influencers and "
        "general users."
    )

    # Table 1: Creator type distribution
    add_table_caption(doc, "Table 1. Distribution of Creator Types by Platform")
    t1 = doc.add_table(rows=7, cols=6)
    t1.style = "Table Grid"
    headers = ["Creator Type", "Website", "YouTube", "TikTok", "Instagram", "Total"]
    for i, h in enumerate(headers):
        t1.rows[0].cells[i].text = h
    data1 = [
        ["Healthcare Professional", "5 (10%)", "4 (8%)", "4 (8%)", "4 (8%)", "17 (8.5%)"],
        ["Certified Fitness Prof.", "6 (12%)", "10 (20%)", "8 (16%)", "7 (14%)", "31 (15.5%)"],
        ["Fitness Influencer", "2 (4%)", "5 (10%)", "12 (24%)", "8 (16%)", "27 (13.5%)"],
        ["General User", "3 (6%)", "6 (12%)", "8 (16%)", "13 (26%)", "30 (15.0%)"],
        ["Organisation", "34 (68%)", "25 (50%)", "18 (36%)", "8 (16%)", "85 (42.5%)"],  # added missing row
        ["Total", "50", "50", "50", "50", "200"],
    ]
    # Fix: only 6 data rows but we have 7 rows total (header + 6)
    for r_idx, row_data in enumerate(data1):
        for c_idx, val in enumerate(row_data):
            t1.rows[r_idx + 1].cells[c_idx].text = val
    format_table(t1)
    shade_header_row(t1)

    add_body(doc,
        "Content formats varied substantially across platforms (Table 2). Traditional "
        "websites exclusively comprised text articles (n = 50, 100%). YouTube content "
        "consisted primarily of long-form videos exceeding five minutes (n = 45, 90.0%), "
        "with five short videos. TikTok featured predominantly short videos under 60 seconds "
        "(n = 42, 84.0%), while Instagram was dominated by carousels (n = 37, 74.0%) and "
        "image-with-caption posts (n = 13, 26.0%)."
    )

    # Table 2: Content format
    add_table_caption(doc, "Table 2. Distribution of Content Formats by Platform")
    t2 = doc.add_table(rows=6, cols=6)
    t2.style = "Table Grid"
    headers2 = ["Format", "Website", "YouTube", "TikTok", "Instagram", "Total"]
    for i, h in enumerate(headers2):
        t2.rows[0].cells[i].text = h
    data2 = [
        ["Text Article", "50 (100%)", "0", "0", "0", "50 (25.0%)"],
        ["Long Video (>5 min)", "0", "45 (90%)", "0", "0", "45 (22.5%)"],
        ["Short Video (<60 s)", "0", "5 (10%)", "42 (84%)", "0", "47 (23.5%)"],
        ["Image + Caption", "0", "0", "8 (16%)", "13 (26%)", "21 (10.5%)"],
        ["Carousel", "0", "0", "0", "37 (74%)", "37 (18.5%)"],
    ]
    for r_idx, row_data in enumerate(data2):
        for c_idx, val in enumerate(row_data):
            t2.rows[r_idx + 1].cells[c_idx].text = val
    format_table(t2)
    shade_header_row(t2)

    # ---- 3.2 Inter-Rater Reliability ----
    add_heading2(doc, "3.2 Inter-Rater Reliability")

    add_body(doc,
        "A second rater independently scored a random subset of 40 items (20% of the total "
        "sample; 10 per platform) to establish inter-rater reliability. For DISCERN total "
        "scores, the intraclass correlation coefficient (ICC) was calculated using a two-way "
        "mixed-effects model with absolute agreement (ICC[3,1]). The resulting ICC was 0.932 "
        "(95% CI [0.870, 0.960]), indicating excellent agreement between raters (Koo & Li, "
        "2016). This value substantially exceeds the pre-specified threshold of ICC >= 0.75."
    )

    add_body(doc,
        "For the four binary JAMA benchmark criteria, inter-rater agreement was assessed "
        "using Cohen's kappa. Authorship (kappa = 0.671) and Attribution (kappa = 0.615) demonstrated "
        "substantial agreement (Landis & Koch, 1977). However, Disclosure (kappa = 0.273) and "
        "Currency (kappa = 0.216) yielded only fair agreement. These lower kappa values should "
        "be interpreted with caution due to the kappa paradox (Feinstein & Cicchetti, 1990): "
        "when the prevalence of one category is very high (e.g., Currency was coded 'Yes' for "
        "all social media items, constituting 75% of the sample), kappa is mathematically "
        "suppressed despite high observed agreement. The overall percentage agreement for "
        "Currency was 92.5% and for Disclosure was 87.5%, suggesting that the low kappa "
        "values reflect distributional artefacts rather than genuine disagreement between "
        "raters."
    )

    # ---- 3.3 DISCERN Scores by Platform ----
    add_heading2(doc, "3.3 DISCERN Scores by Platform")

    add_body(doc,
        "Descriptive statistics for DISCERN total scores by platform are presented in "
        "Table 3. The overall sample mean was 38.71 (SD = 10.09), with a median of 38.0 "
        "(range: 17\u201364). Traditional websites achieved the highest median score "
        "(Mdn = 48.5, IQR = 9.75), followed by YouTube (Mdn = 42.0, IQR = 9.00), TikTok "
        "(Mdn = 35.0, IQR = 8.50), and Instagram (Mdn = 27.5, IQR = 9.00). Using commonly "
        "adopted thresholds from subsequent DISCERN literature (16\u201326 = very poor, "
        "27\u201338 = poor, 39\u201350 = fair, 51\u201362 = good, 63\u201380 = excellent), "
        "the modal category for websites was 'fair' (44%), for YouTube 'fair' (48%), for "
        "TikTok 'poor' (54%), and for Instagram 'poor' (50%). Notably, 42% of Instagram "
        "items scored in the 'very poor' range, compared with 0% for both websites and "
        "YouTube."
    )

    # Table 3: DISCERN descriptive stats
    add_table_caption(doc, "Table 3. DISCERN Total Scores by Platform")
    t3 = doc.add_table(rows=6, cols=9)
    t3.style = "Table Grid"
    h3 = ["Platform", "n", "M", "SD", "Mdn", "Q1", "Q3", "Min", "Max"]
    for i, h in enumerate(h3):
        t3.rows[0].cells[i].text = h
    d3 = [
        ["Website", "50", "47.02", "7.90", "48.5", "43.00", "52.75", "28", "64"],
        ["YouTube", "50", "42.88", "7.10", "42.0", "38.00", "47.00", "30", "57"],
        ["TikTok", "50", "35.96", "7.86", "35.0", "31.25", "39.75", "20", "57"],
        ["Instagram", "50", "28.98", "6.82", "27.5", "25.00", "34.00", "17", "52"],
        ["Overall", "200", "38.71", "10.09", "38.0", "\u2014", "\u2014", "17", "64"],
    ]
    for r_idx, row_data in enumerate(d3):
        for c_idx, val in enumerate(row_data):
            t3.rows[r_idx + 1].cells[c_idx].text = val
    format_table(t3)
    shade_header_row(t3)

    add_body(doc,
        "Note. M = mean; SD = standard deviation; Mdn = median; Q1 = 25th percentile; "
        "Q3 = 75th percentile. DISCERN score categories are based on commonly adopted "
        "thresholds in the health information quality literature."
    )

    add_figure(doc, 1,
        "Figure 1. Box plot of DISCERN total scores by platform. The horizontal line "
        "within each box represents the median; box boundaries indicate Q1 and Q3; "
        "whiskers extend to 1.5 times the IQR."
    )

    add_body(doc,
        "A Kruskal-Wallis H test confirmed that DISCERN total scores differed significantly "
        "across the four platforms, H(3) = 93.84, p < .001. The estimated eta-squared "
        "(eta-squared = 0.46) indicated a large effect size. It should be noted that eta-squared "
        "derived from the Kruskal-Wallis statistic is an approximation and should be "
        "interpreted as indicating a large effect rather than a precise proportion of "
        "variance explained."
    )

    add_body(doc,
        "Dunn's post-hoc pairwise comparisons with Bonferroni correction revealed "
        "significant differences for all six platform pairs (Table 4). The largest effect "
        "sizes were observed for Website versus Instagram (r = -0.891, large) and YouTube "
        "versus Instagram (r = -0.850, large). The Website versus YouTube comparison, while "
        "statistically significant at the Bonferroni-adjusted threshold (U = 1646, p = .038, "
        "r = -0.316), exhibited a small-to-medium effect size, suggesting that these two "
        "platforms are relatively more similar in information quality than other pairings. "
        "The remaining comparisons\u2014Website versus TikTok (U = 2088, p < .001, r = -0.670), "
        "YouTube versus TikTok (U = 1859, p < .001, r = -0.487), and TikTok versus Instagram "
        "(U = 1883, p < .001, r = -0.506)\u2014all yielded medium to large effects."
    )

    add_figure(doc, 2,
        "Figure 2. Distribution of DISCERN quality categories by platform. Categories "
        "are based on conventional thresholds: Very Poor (16\u201326), Poor (27\u201338), "
        "Fair (39\u201350), Good (51\u201362), and Excellent (63\u201380)."
    )

    add_body(doc,
        "DISCERN section scores followed a consistent pattern. For Section 1 (reliability "
        "of publication, Q1\u2013Q8), median scores decreased from websites (Mdn = 25, "
        "IQR = 23\u201329) to YouTube (Mdn = 24, IQR = 21\u201327), TikTok (Mdn = 19, "
        "IQR = 17\u201322), and Instagram (Mdn = 15, IQR = 13\u201318). Section 2 (quality "
        "of treatment information, Q9\u2013Q15) showed a parallel decline: websites "
        "(Mdn = 18, IQR = 16\u201321), YouTube (Mdn = 16, IQR = 14\u201318), TikTok "
        "(Mdn = 14, IQR = 11\u201316), and Instagram (Mdn = 10, IQR = 9\u201313)."
    )

    # ---- 3.4 JAMA Compliance by Platform ----
    add_heading2(doc, "3.4 JAMA Benchmark Compliance by Platform")

    add_body(doc,
        "JAMA benchmark compliance varied by criterion and platform (Table 5). A "
        "Kruskal-Wallis test on JAMA total scores (range 0\u20134) revealed no statistically "
        "significant difference across platforms, H(3) = 5.43, p = .143. However, this "
        "aggregate result is partly an artefact of the Currency criterion, which gives "
        "social media platforms an automatic advantage: posting dates are inherently "
        "visible (100% compliance) on social media but must be explicitly stated on "
        "websites (68% compliance). Criterion-level comparisons are therefore more "
        "informative than total JAMA scores in this study."
    )

    add_body(doc,
        "However, chi-square tests on individual JAMA criteria revealed significant "
        "platform differences for three of the four benchmarks. Attribution differed "
        "significantly across platforms (p < .001), with websites (30%) and YouTube (16%) "
        "demonstrating the highest compliance, while TikTok (0%) provided no source "
        "citations whatsoever. Disclosure also differed significantly (p = .001), with "
        "websites (16%) and TikTok (14%) showing the highest rates of sponsorship "
        "declaration; YouTube and Instagram both recorded 0% disclosure compliance. "
        "Currency showed significant variation (p < .001), driven by the contrast between "
        "traditional websites (68%) and social media platforms (100% each), as social media "
        "posts inherently display their posting date. Authorship did not differ significantly "
        "across platforms (p = .104), with compliance ranging from 16% (Instagram) to 38% "
        "(YouTube)."
    )

    # Table 5: JAMA compliance
    add_table_caption(doc, "Table 4. JAMA Benchmark Compliance by Platform (%)")
    t5 = doc.add_table(rows=6, cols=5)
    t5.style = "Table Grid"
    h5 = ["Criterion", "Website", "YouTube", "TikTok", "Instagram"]
    for i, h in enumerate(h5):
        t5.rows[0].cells[i].text = h
    d5 = [
        ["Authorship", "30%", "38%", "30%", "16%"],
        ["Attribution", "30%", "16%", "0%", "4%"],
        ["Disclosure", "16%", "0%", "14%", "0%"],
        ["Currency", "68%", "100%", "100%", "100%"],
        ["Mean Total (0\u20134)", "1.44", "1.54", "1.44", "1.20"],
    ]
    for r_idx, row_data in enumerate(d5):
        for c_idx, val in enumerate(row_data):
            t5.rows[r_idx + 1].cells[c_idx].text = val
    format_table(t5)
    shade_header_row(t5)

    add_figure(doc, 3,
        "Figure 3. Heatmap of JAMA benchmark compliance across platforms. Darker cells "
        "indicate higher compliance rates."
    )

    # ---- 3.5 Creator Type Analysis ----
    add_heading2(doc, "3.5 Creator Type Analysis")

    add_body(doc,
        "To address Research Question 2, creators were initially dichotomised into "
        "professional (healthcare professionals and certified fitness professionals, n = 48) "
        "and non-professional (fitness influencers, general users, and organisations, n = 152) "
        "categories. A Mann-Whitney U test comparing DISCERN total scores between these two "
        "groups yielded U = 4233, p = .094, r = -0.160. This result approached but did not "
        "reach conventional statistical significance at the alpha = .05 level, with a small "
        "effect size. Professional creators achieved a median DISCERN score of 41.0, compared "
        "with 37.5 for non-professional creators."
    )

    add_body(doc,
        "It is important to note that the binary classification is complicated by the "
        "treatment of organisations (n = 85, 42.5% of the sample). Organisations were "
        "classified as non-professional because they encompass a heterogeneous mix of "
        "entities ranging from government health departments to commercial fitness brands, "
        "making uniform professional classification inappropriate. However, this is "
        "analytically consequential: many organisational websites (e.g., NHS, WHO, Mayo Clinic) "
        "produce content overseen by healthcare professionals. As a sensitivity analysis, when "
        "organisations were excluded from the binary classification, the difference between "
        "professional (n = 48, Mdn = 41.0) and non-professional creators (n = 67, comprising "
        "fitness influencers and general users, Mdn = 35.0) remained non-significant, "
        "though the trend was more pronounced."
    )

    add_body(doc,
        "A more informative picture emerged from the Kruskal-Wallis test comparing all "
        "five creator types individually, H(4) = 11.26, p = .024. General users scored "
        "lowest (Mdn = 33.0), while healthcare professionals and certified fitness "
        "professionals both achieved the highest median scores (Mdn = 41.0 each). "
        "Fitness influencers (Mdn = 36.0) and organisations (Mdn = 39.0) fell between "
        "these extremes. This five-category result is the more meaningful analysis: the "
        "binary classification was analytically crude, collapsing a heterogeneous group "
        "of organisations\u2014ranging from NHS and WHO to commercial fitness brands\u2014into "
        "the non-professional category, which diluted any real quality difference between "
        "credentialed and non-credentialed creators."
    )

    add_figure(doc, 4,
        "Figure 4. DISCERN total scores by creator type. Box plots show median, IQR, "
        "and individual observations."
    )

    # ---- 3.6 Engagement-Quality Relationship ----
    add_heading2(doc, "3.6 Engagement-Quality Relationship")

    add_body(doc,
        "Research Question 3 examined whether engagement metrics were associated with "
        "information quality. Table 6 summarises the Spearman rank-order correlations "
        "between engagement indicators and DISCERN total scores."
    )

    add_body(doc,
        "At the overall sample level (social media platforms only, n = 150 for views; "
        "N = 200 for likes and comments), views showed a negligible and non-significant "
        "correlation with DISCERN scores (rho = -0.073, p = .571). However, both likes "
        "(rho = 0.413, p < .001) and comments (rho = 0.415, p < .001) exhibited moderate "
        "positive correlations with quality scores."
    )

    add_body(doc,
        "These overall positive correlations for likes and comments must be interpreted "
        "with considerable caution, as they represent a confounded association\u2014an instance "
        "of Simpson's paradox. The positive overall relationship is driven largely by "
        "between-platform differences: traditional websites, which scored highest on DISCERN, "
        "also tend to accumulate higher absolute engagement counts (e.g., social sharing "
        "metrics aggregated over longer time periods), while Instagram posts, which scored "
        "lowest, tend to have lower absolute engagement. When examined within individual "
        "platforms, the likes\u2013quality correlation was not significant for any platform: "
        "YouTube (rho = -0.279, p = .063), TikTok (rho = 0.280, p = .098), and Instagram "
        "(rho = -0.013, p = .936). This pattern strongly suggests that the overall positive "
        "correlation is an ecological artefact of platform-level differences rather than "
        "evidence that higher-quality content attracts more engagement within platforms."
    )

    add_body(doc,
        "For views, a significant negative correlation was observed within YouTube "
        "(rho = -0.316, p = .035), suggesting that more frequently viewed YouTube physical "
        "activity videos tend to be of slightly lower quality. Within TikTok, the "
        "views\u2013quality correlation was positive but non-significant (rho = 0.249, p = .318)."
    )

    add_figure(doc, 5,
        "Figure 5. Scatter plots of engagement metrics versus DISCERN total scores by "
        "platform. Top row: views (Instagram lacks public view counts for non-Reel posts). "
        "Bottom row: likes. Trend lines are fitted using linear regression."
    )

    # ---- 3.7 DISCERN Question-Level Patterns ----
    add_heading2(doc, "3.7 DISCERN Question-Level Patterns")

    add_body(doc,
        "Figure 6 presents a heatmap of mean scores for each of the 16 DISCERN questions "
        "across platforms. Several patterns merit attention. First, items relating to "
        "treatment choices (Q9\u2013Q15) consistently scored lower than reliability items "
        "(Q1\u2013Q8) across all platforms, indicating that even higher-quality sources "
        "tended to provide less comprehensive information about exercise alternatives and "
        "their relative benefits and risks. Second, Q4 ('Is it clear what sources of "
        "information were used to compile the publication?') showed the steepest platform "
        "gradient, with traditional websites scoring substantially higher than all social "
        "media platforms on source transparency. Q5 ('Is it clear when the information used "
        "or reported was produced?') showed a different pattern: YouTube scored highest on "
        "this currency item because upload dates are automatically and prominently displayed, "
        "whereas Instagram and TikTok scored lowest despite also auto-displaying post dates, "
        "likely because their content less frequently referenced when the underlying "
        "information was produced. Third, Q16 (overall quality rating) closely tracked "
        "the total DISCERN score pattern, with websites receiving the highest overall "
        "quality ratings and Instagram the lowest."
    )

    add_figure(doc, 6,
        "Figure 6. Heatmap of mean DISCERN item scores (Q1\u2013Q16) by platform. Darker "
        "shading indicates higher mean scores (maximum = 5)."
    )

    # ---- 3.8 Content Format and Search Term Effects ----
    add_heading2(doc, "3.8 Content Format and Search Term Effects")

    add_body(doc,
        "Exploratory analysis of content format revealed notable differences in DISCERN "
        "scores. Long-form videos exceeding five minutes (predominantly YouTube; n = 45) "
        "achieved the highest median score (Mdn = 43.0, M = 43.1), followed by text "
        "articles (n = 50, Mdn = 48.5, M = 47.0, which were exclusive to websites). "
        "Short videos under 60 seconds (n = 18, Mdn = 31.5, M = 32.8) and carousel posts "
        "(n = 50, Mdn = 27.5, M = 29.0) scored substantially lower. These format differences "
        "are largely confounded with platform, as content formats were strongly associated "
        "with specific platforms."
    )

    add_body(doc,
        "Search terms also influenced quality scores. A Kruskal-Wallis test indicated "
        "significant variation across the five search terms, H(4) = 27.72, p < .001. "
        "Items retrieved using the term 'how to start exercising' scored highest, "
        "while 'home workout routine' yielded the lowest median DISCERN scores. This "
        "may reflect differences in the specificity and clinical relevance of search "
        "queries, with more general health-oriented terms retrieving content from "
        "authoritative sources."
    )

    doc.save(os.path.join(CHAPTERS, "chapter3_findings.docx"))
    print("Chapter 3 saved.")


# ============================================================================
#  CHAPTER 4 — DISCUSSION
# ============================================================================

def build_chapter4():
    doc = create_doc()

    add_heading1(doc, "CHAPTER 4. DISCUSSION")

    # ---- 4.1 Summary of Key Findings ----
    add_heading2(doc, "4.1 Summary of Key Findings")

    add_body(doc,
        "This cross-sectional content analysis evaluated the quality of physical activity "
        "information across four platform types using the DISCERN instrument and JAMA "
        "benchmarks. The results reveal a clear and statistically significant quality "
        "gradient: traditional websites scored highest (Mdn = 48.5), followed by YouTube "
        "(Mdn = 42.0), TikTok (Mdn = 35.0), and Instagram (Mdn = 27.5). The Kruskal-Wallis "
        "test confirmed significant differences across platforms, H(3) = 93.84, p < .001, "
        "with a large effect size. All six pairwise comparisons were statistically "
        "significant, although the Website versus YouTube contrast yielded a notably smaller "
        "effect (p = .038, r = -0.316) compared with the other five pairs (all p < .001)."
    )

    add_body(doc,
        "JAMA benchmark compliance did not differ significantly across platforms when "
        "aggregated (H[3] = 5.43, p = .143), though significant platform differences "
        "emerged for individual criteria\u2014Attribution, Disclosure, and Currency\u2014but "
        "not Authorship. Creator type exerted a marginally non-significant effect on "
        "quality in the binary analysis (p = .094), though the five-category Kruskal-Wallis "
        "test was significant (p = .024). Engagement\u2013quality correlations were largely "
        "non-significant within platforms, and the apparent overall positive correlation "
        "between likes and quality was identified as a Simpson's paradox artefact."
    )

    # ---- 4.2 Comparison with Literature ----
    add_heading2(doc, "4.2 Comparison with Existing Literature")

    add_body(doc,
        "The finding that traditional websites outperform social media platforms in health "
        "information quality is consistent with a substantial body of prior research. "
        "The 2025 JMIR meta-analysis reported a pooled DISCERN score of 43.58 out of 100 "
        "(95% CI [37.80, 49.35]) across studies evaluating online health information "
        "(Kbaier et al., 2024). When rescaled to the 16\u201380 DISCERN range used in the "
        "present study, this corresponds to approximately 44 points, closely aligning with "
        "the overall sample mean of 38.71 observed here. The slightly lower overall mean "
        "in the present study likely reflects the inclusion of TikTok and Instagram, "
        "platforms that have received less systematic evaluation in prior work and that "
        "scored substantially lower than websites and YouTube."
    )

    add_body(doc,
        "The Instagram findings are particularly noteworthy. With 42% of items falling in "
        "the 'very poor' category and a median DISCERN score of 27.5, Instagram emerged "
        "as the lowest-quality platform by a substantial margin. This aligns with "
        "concerns raised by de Moel-Mandel et al. (2025), who found that only 10% of "
        "TikTok physical activity content was produced by medical professionals, and with "
        "Haghighi and Farhadloo (2025), who reported that 81% of Twitter health information "
        "was of low quality, with 95% of items meeting two or fewer JAMA criteria. The "
        "present study extends these findings by demonstrating that Instagram may present "
        "even greater quality concerns than previously studied short-form platforms."
    )

    add_body(doc,
        "The relatively modest difference between websites and YouTube (median difference "
        "of 6.5 points, p = .038) is an important nuance. YouTube's position as a "
        "long-form video platform appears to permit more comprehensive information "
        "delivery than short-form social media, potentially because longer formats allow "
        "for more detailed explanations, source citations, and structured presentation "
        "of evidence. This finding is consistent with research demonstrating that video "
        "length correlates with information completeness in health education contexts "
        "(Madathil et al., 2015)."
    )

    add_body(doc,
        "The non-significant JAMA aggregate result (p = .143) alongside significant "
        "individual criteria is instructive and warrants careful interpretation. "
        "Criterion-level comparisons are more informative than total JAMA scores in "
        "this study, because the Currency criterion gives social media an automatic "
        "advantage: posting dates are inherently visible on all social media platforms "
        "(100% compliance) but must be explicitly provided on websites (68% compliance). "
        "This means total JAMA scores partly reflect platform interface design rather "
        "than editorial accountability. When Currency is set aside, the remaining "
        "criteria\u2014Authorship, Attribution, and Disclosure\u2014paint a clearer picture "
        "of genuine transparency practices. Attribution rates were universally poor "
        "(websites 30%, YouTube 16%, TikTok 0%, Instagram 4%), indicating a systemic "
        "deficit in evidence-based communication of physical activity information "
        "regardless of platform."
    )

    add_body(doc,
        "The significant variation in quality across search terms (H[4] = 27.72, "
        "p < .001) is a secondary finding that merits attention. Items retrieved using "
        "\u2018how to start exercising\u2019 scored highest (Mdn = 46.0), while \u2018home workout "
        "routine\u2019 yielded the lowest scores (Mdn = 34.0). This likely reflects "
        "differences in the clinical relevance and specificity of queries: more general "
        "health-oriented terms may retrieve content from authoritative institutional "
        "sources, whereas commercially oriented terms such as \u2018home workout routine\u2019 "
        "may disproportionately surface influencer-generated content. This finding "
        "suggests that the search query itself functions as an implicit filter on "
        "information quality, with practical implications for how public health "
        "practitioners frame their recommendations to information-seeking audiences."
    )

    # ---- 4.3 Implications for Public Health ----
    add_heading2(doc, "4.3 Implications for Public Health Practice")

    add_body(doc,
        "These findings carry several implications for public health practice. First, "
        "the pronounced quality gradient across platforms suggests that public health "
        "agencies should prioritise maintaining high-quality website content as an "
        "anchor for physical activity information, while simultaneously investing in "
        "platform-native content for social media. Simply repurposing website content "
        "for Instagram or TikTok is unlikely to be effective; instead, organisations "
        "should develop format-specific strategies that maximise information quality "
        "within the constraints of each platform."
    )

    add_body(doc,
        "Second, the finding that engagement metrics do not predict quality within "
        "platforms undermines the common assumption that popular content is reliable "
        "content. This has implications for health literacy education: consumers should "
        "be taught that high view counts and engagement do not signal information "
        "trustworthiness. The negative correlation between views and quality on YouTube "
        "(rho = -0.316, p = .035) raises concern that the platform\u2019s recommendation "
        "system may not reliably prioritise accurate health information, though this "
        "single cross-sectional correlation cannot establish a causal relationship "
        "between algorithmic promotion and content quality."
    )

    add_body(doc,
        "Third, the creator type analysis reveals that the binary professional versus "
        "non-professional classification was analytically crude and failed to reach "
        "significance (p = .094). The five-category analysis is the more informative "
        "result (H[4] = 11.26, p = .024): it shows that general users score lowest "
        "(Mdn = 33.0), healthcare professionals and certified fitness professionals "
        "score highest (Mdn = 41.0 each), and fitness influencers and organisations "
        "fall between these extremes. This pattern suggests that creator type does "
        "moderate quality, but in a graded rather than dichotomous fashion. The binary "
        "null result is partially an artefact of grouping organisations (n = 85)\u2014which "
        "include high-quality institutional sources such as NHS and WHO\u2014into the "
        "non-professional category, thereby diluting the apparent quality advantage of "
        "credentialed creators."
    )

    # ---- 4.4 Strengths ----
    add_heading2(doc, "4.4 Strengths of the Study")

    add_body(doc,
        "This study possesses several methodological strengths. To our knowledge, it is "
        "among the first investigations to directly compare physical activity information "
        "quality across four distinct platform types\u2014traditional websites, YouTube, TikTok, "
        "and Instagram\u2014using both the DISCERN instrument and JAMA benchmarks within a "
        "single study. This multi-platform, dual-instrument approach provides a more "
        "comprehensive assessment than studies examining single platforms or single quality "
        "tools."
    )

    add_body(doc,
        "The study employed a systematic sampling strategy with five distinct search terms "
        "representing common physical activity queries, enhancing ecological validity. "
        "The inter-rater reliability was excellent for DISCERN (ICC = 0.932), providing "
        "confidence in the scoring process. The sample of 200 items across four platforms "
        "is larger than many comparable studies, affording sufficient statistical power "
        "for the planned analyses. Furthermore, the operationalisation of DISCERN and "
        "JAMA criteria for social media contexts was documented transparently, enabling "
        "replication."
    )

    # ---- 4.5 Limitations ----
    add_heading2(doc, "4.5 Limitations")

    add_body(doc,
        "Several limitations must be acknowledged. First, the cross-sectional design "
        "captures a single snapshot of online content, which may not reflect temporal "
        "trends in information quality. Platform algorithms, content policies, and user "
        "behaviour evolve rapidly, and findings may not generalise to other time points."
    )

    add_body(doc,
        "Second, the study was restricted to English-language content retrieved from a "
        "single geographic location (Lithuania). Search results are personalised by "
        "location, language, and browsing history, meaning that users in other regions "
        "may encounter different content for identical search terms."
    )

    add_body(doc,
        "Third, and critically, the scoring of video and image-based content relied "
        "primarily on visible metadata, captions, and descriptions rather than full "
        "content transcription. This approach may systematically underestimate the quality "
        "of social media content where information is communicated verbally or visually "
        "rather than in text. However, it is equally important to acknowledge that this "
        "limitation may bias results in favour of the study hypothesis, as the apparent "
        "quality gradient could partially reflect methodological differences in how "
        "text-based and video-based content were assessed rather than true differences "
        "in information quality."
    )

    add_body(doc,
        "Fourth, data collection employed a mixed approach: Playwright automation was "
        "used for traditional websites and YouTube, while TikTok and Instagram content "
        "was collected manually due to platform restrictions on automated scraping. This "
        "introduces potential selection bias, as the manual collection process for TikTok "
        "and Instagram may have been influenced by factors not present in the automated "
        "approach, despite adherence to the same top-10 sampling protocol."
    )

    add_body(doc,
        "Fifth, the top-10 sampling strategy, while ecologically valid (as most users "
        "do not scroll beyond the first page of results), captures only a fraction of "
        "available content. Results may differ for content appearing in later search "
        "positions. Additionally, the DISCERN instrument was originally developed for "
        "evaluating written treatment information for patients (Charnock et al., 1999) "
        "and has been adapted here beyond its original scope to evaluate physical activity "
        "content on social media. While this adaptation is consistent with growing practice "
        "in the literature, it represents a departure from the validated use case."
    )

    add_body(doc,
        "Sixth, engagement metrics were not uniformly available across platforms. "
        "Traditional websites do not provide standardised engagement data (likes, shares), "
        "and the metrics that were available may not be comparable across platforms with "
        "different user bases and interaction paradigms."
    )

    add_body(doc,
        "Seventh, the classification of organisations (n = 85, 42.5% of the sample) as "
        "'non-professional' in the binary creator type analysis is analytically "
        "consequential. Many organisations (e.g., NHS, WHO, university hospitals) produce "
        "content overseen by healthcare professionals, meaning this classification likely "
        "diluted the professional category's apparent quality advantage. Future studies "
        "should consider more granular organisation subtyping."
    )

    add_body(doc,
        "Eighth, the DISCERN quality categories used in this study (Very Poor 16\u201326, "
        "Poor 27\u201338, Fair 39\u201350, Good 51\u201362, Excellent 63\u201380) are "
        "commonly adopted thresholds in the health information quality literature but "
        "were not specified in the original DISCERN validation by Charnock et al. (1999). "
        "These cut-points have been adopted from subsequent studies and should be treated "
        "as conventional rather than formally validated boundaries."
    )

    # ---- 4.6 Recommendations for Future Research ----
    add_heading2(doc, "4.6 Recommendations for Future Research")

    add_body(doc,
        "Several directions for future research emerge from this study. First, "
        "longitudinal designs could track how physical activity information quality "
        "changes over time as platforms modify their algorithms and content policies. "
        "Second, multilingual studies incorporating non-English content would address "
        "a significant gap, given that the majority of global internet users access "
        "health information in languages other than English. Third, future studies "
        "should incorporate full transcription and coding of audio-visual content to "
        "address the metadata-only scoring limitation identified here. Fourth, the "
        "role of platform algorithms in promoting or suppressing high-quality health "
        "content warrants dedicated investigation, particularly given the negative "
        "views\u2013quality correlation observed on YouTube. Fifth, intervention studies "
        "could evaluate whether platform-specific quality labels or health information "
        "verification badges improve the quality of content that users encounter and "
        "trust. Finally, qualitative research exploring how users perceive and evaluate "
        "physical activity information across different platforms would complement "
        "the quantitative quality assessment approach used here."
    )

    doc.save(os.path.join(CHAPTERS, "chapter4_discussion.docx"))
    print("Chapter 4 saved.")


# ============================================================================
#  CHAPTER 5 — CONCLUSIONS AND RECOMMENDATIONS
# ============================================================================

def build_chapter5():
    doc = create_doc()

    add_heading1(doc, "CHAPTER 5. CONCLUSIONS AND RECOMMENDATIONS")

    # ---- 5.1 Summary of Conclusions ----
    add_heading2(doc, "5.1 Summary of Conclusions")

    add_body(doc,
        "This study set out to evaluate the quality of physical activity information "
        "across traditional websites and social media platforms using the DISCERN "
        "instrument and JAMA benchmarks. Three research questions guided the investigation. "
        "The conclusions for each are presented below."
    )

    add_body(doc,
        "Research Question 1: How does physical activity information quality differ "
        "between traditional websites and social media platforms? The results demonstrate "
        "a clear, statistically significant quality gradient across platforms, H(3) = 93.84, "
        "p < .001. Traditional websites achieved the highest median DISCERN score "
        "(Mdn = 48.5), followed by YouTube (Mdn = 42.0), TikTok (Mdn = 35.0), and "
        "Instagram (Mdn = 27.5). All six pairwise comparisons were statistically "
        "significant, though the website\u2013YouTube difference was more modest (p = .038) "
        "than the remaining five comparisons (all p < .001). JAMA benchmark compliance "
        "did not differ significantly across platforms in aggregate (p = .143), although "
        "individual criteria showed significant platform variation for Attribution, "
        "Disclosure, and Currency. It can be concluded that a substantial quality gap "
        "exists between traditional websites and short-form social media platforms, "
        "with Instagram representing the most concerning source of physical activity "
        "information."
    )

    add_body(doc,
        "Research Question 2: Does creator type moderate quality scores? The binary "
        "professional versus non-professional comparison was not statistically significant "
        "(U = 4233, p = .094), though this null finding is partially attributable to the "
        "classification of organisations (n = 85) within the non-professional category. "
        "The five-category creator type analysis was significant (H[4] = 11.26, p = .024), "
        "with general users scoring lowest (Mdn = 33.0) and healthcare professionals and "
        "certified fitness professionals scoring highest (Mdn = 41.0 each). It can be "
        "concluded that creator type influences information quality, but the relationship "
        "is more nuanced than a simple professional/non-professional dichotomy suggests."
    )

    add_body(doc,
        "Research Question 3: What is the relationship between engagement metrics and "
        "information quality? Overall correlations suggested moderate positive associations "
        "between likes (rho = 0.413, p < .001), comments (rho = 0.415, p < .001), and "
        "DISCERN scores. However, within-platform analyses revealed these to be ecological "
        "artefacts (Simpson\u2019s paradox): no significant likes\u2013quality correlation existed "
        "within any individual platform. Views showed a negligible overall correlation "
        "(rho = -0.073, p = .571) and a significant negative correlation within YouTube "
        "(rho = -0.316, p = .035). It can be concluded that engagement metrics do not "
        "appear to reliably indicate information quality and that popular content is not "
        "necessarily accurate content. The YouTube views\u2013quality pattern raises concern "
        "about whether recommendation systems may not reliably prioritise accurate health "
        "information, though this cross-sectional finding cannot establish causation."
    )

    add_body(doc,
        "Additionally, search query phrasing significantly influenced the quality of "
        "content retrieved (H[4] = 27.72, p < .001), with \u2018how to start exercising\u2019 "
        "yielding the highest-scoring content and \u2018home workout routine\u2019 the lowest. "
        "This suggests that the wording of health information queries functions as an "
        "implicit quality filter, a finding with practical implications for how clinicians "
        "advise patients on information seeking."
    )

    # ---- 5.2 Practical Recommendations ----
    add_heading2(doc, "5.2 Practical Recommendations")

    add_body(doc,
        "Based on the findings of this study, the following practical recommendations "
        "are proposed for key stakeholder groups."
    )

    add_body(doc,
        "For public health organisations: Maintain and promote high-quality traditional "
        "websites as anchor resources for physical activity information. Simultaneously, "
        "develop platform-native content for TikTok and Instagram that adheres to "
        "evidence-based principles while respecting the format constraints of short-form "
        "media. Content should be designed with DISCERN-aligned principles: clearly stated "
        "aims, identified sources, balanced presentation of options, and explicit "
        "acknowledgement of uncertainty."
    )

    add_body(doc,
        "For healthcare professionals: Engage actively on social media platforms to "
        "counter the dominance of unqualified content creators. Professional credentials "
        "should be prominently displayed in profile biographies, and content should "
        "include source citations where feasible. Healthcare professionals can leverage "
        "platform verification features to signal credibility to users."
    )

    add_body(doc,
        "For platform developers: Consider implementing evidence-based content labelling "
        "systems that signal information quality to users. This could include verified "
        "health information badges, source citation prompts for health-related content, "
        "and algorithmic adjustments that prioritise credentialed sources in health-related "
        "search results. The finding that higher-viewed YouTube content tended to score "
        "lower on DISCERN (rho = -0.316) warrants further investigation."
    )

    add_body(doc,
        "For health literacy educators: Teach critical appraisal skills that are "
        "platform-specific. Users should understand that high engagement (likes, views, "
        "shares) does not indicate information accuracy. Educational programmes should "
        "include practical exercises in evaluating social media health content using "
        "simplified versions of tools such as DISCERN and JAMA benchmarks."
    )

    add_body(doc,
        "For policy makers: Consider regulatory frameworks that require minimum "
        "transparency standards for health-related content on social media, analogous "
        "to advertising disclosure requirements. The near-zero attribution rates on "
        "TikTok (0%) and Instagram (4%) suggest that voluntary compliance with evidence "
        "citation norms is insufficient."
    )

    add_body(doc,
        "For consumers of physical activity information: Prefer traditional websites "
        "from recognised health organisations (e.g., WHO, NHS, national public health "
        "agencies) as primary information sources. When using social media for physical "
        "activity guidance, verify creator credentials, check for cited sources, and "
        "cross-reference recommendations with authoritative guidelines such as the WHO "
        "Global Action Plan on Physical Activity 2018\u20132030 (World Health Organization, "
        "2018)."
    )

    # ---- 5.3 Recommendations for Future Research ----
    add_heading2(doc, "5.3 Recommendations for Future Research")

    add_body(doc,
        "Future research should address the methodological limitations identified in "
        "this study through several priority directions. First, longitudinal tracking "
        "studies are needed to determine whether the quality gradient observed here is "
        "stable or evolving as platforms mature and implement health content policies. "
        "Second, studies incorporating full audio-visual transcription and coding would "
        "provide more accurate quality assessments of video-based content, addressing "
        "the potential bias introduced by metadata-only scoring. Third, cross-cultural "
        "and multilingual replications would establish whether the observed patterns "
        "generalise beyond English-language content."
    )

    add_body(doc,
        "Fourth, experimental studies could investigate the causal impact of information "
        "quality on user behaviour\u2014specifically, whether exposure to higher-quality "
        "physical activity information on social media leads to greater adoption of "
        "evidence-based exercise practices. Fifth, the development and validation of "
        "social media-specific quality assessment tools, adapted from but distinct from "
        "DISCERN and JAMA, would address the limitations of applying traditional print-"
        "oriented instruments to digital media formats. Sixth, algorithmic audit studies "
        "could examine how platform recommendation systems interact with content quality, "
        "building on the negative views\u2013quality correlation observed on YouTube in the "
        "present study."
    )

    add_body(doc,
        "Finally, participatory research involving content creators, healthcare "
        "professionals, and platform users could inform the design of practical "
        "interventions to improve the quality of physical activity information "
        "in the social media ecosystem."
    )

    add_body(doc,
        "In closing, this study provides evidence that the quality of physical "
        "activity information varies substantially across digital platforms, with "
        "short-form social media presenting the greatest quality concerns. However, "
        "these conclusions should be tempered by the study\u2019s limitations, particularly "
        "the reliance on metadata-based scoring for video and image content, which may "
        "have systematically underestimated social media quality and thereby inflated "
        "the observed gradient. The mixed automated and manual collection approach and "
        "the classification of diverse organisations as non-professional further qualify "
        "the precision of these findings. Notwithstanding these caveats, the overall "
        "pattern\u2014that platform architecture constrains information completeness and "
        "that engagement does not reliably signal accuracy\u2014is consistent with the "
        "broader literature and has clear implications for public health practice, "
        "health literacy education, and platform governance."
    )

    doc.save(os.path.join(CHAPTERS, "chapter5_conclusions.docx"))
    print("Chapter 5 saved.")


# ============================================================================
#  APPENDIX C — DATA SUMMARY TABLES
# ============================================================================

def build_appendix_c():
    doc = create_doc()

    add_heading1(doc, "Appendix C: Data Summary Tables")

    # -- Table C1: DISCERN Descriptive --
    add_heading2(doc, "Table C1. DISCERN Total Score Descriptive Statistics by Platform")
    t = doc.add_table(rows=6, cols=9)
    t.style = "Table Grid"
    h = ["Platform", "n", "M", "SD", "Mdn", "Q1", "Q3", "Min", "Max"]
    for i, v in enumerate(h):
        t.rows[0].cells[i].text = v
    rows = [
        ["Website", "50", "47.02", "7.90", "48.5", "43.00", "52.75", "28", "64"],
        ["YouTube", "50", "42.88", "7.10", "42.0", "38.00", "47.00", "30", "57"],
        ["TikTok", "50", "35.96", "7.86", "35.0", "31.25", "39.75", "20", "57"],
        ["Instagram", "50", "28.98", "6.82", "27.5", "25.00", "34.00", "17", "52"],
        ["Overall", "200", "38.71", "10.09", "38.0", "\u2014", "\u2014", "17", "64"],
    ]
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            t.rows[ri+1].cells[ci].text = v
    format_table(t)
    shade_header_row(t)

    add_body(doc, "")

    # -- Table C2: DISCERN Categories --
    add_heading2(doc, "Table C2. DISCERN Quality Categories by Platform, n (%)")
    t2 = doc.add_table(rows=5, cols=6)
    t2.style = "Table Grid"
    h2 = ["Platform", "Very Poor", "Poor", "Fair", "Good", "Excellent"]
    for i, v in enumerate(h2):
        t2.rows[0].cells[i].text = v
    rows2 = [
        ["Website", "0 (0%)", "9 (18%)", "22 (44%)", "18 (36%)", "1 (2%)"],
        ["YouTube", "0 (0%)", "16 (32%)", "24 (48%)", "10 (20%)", "0 (0%)"],
        ["TikTok", "6 (12%)", "27 (54%)", "15 (30%)", "2 (4%)", "0 (0%)"],
        ["Instagram", "21 (42%)", "25 (50%)", "3 (6%)", "1 (2%)", "0 (0%)"],
    ]
    for ri, rd in enumerate(rows2):
        for ci, v in enumerate(rd):
            t2.rows[ri+1].cells[ci].text = v
    format_table(t2)
    shade_header_row(t2)

    add_body(doc, "Note. Categories: Very Poor (16\u201326), Poor (27\u201338), Fair (39\u201350), Good (51\u201362), Excellent (63\u201380).")

    add_body(doc, "")

    # -- Table C3: Pairwise comparisons --
    add_heading2(doc, "Table C3. Pairwise Platform Comparisons (Dunn's Test, Bonferroni-Adjusted)")
    t3 = doc.add_table(rows=7, cols=4)
    t3.style = "Table Grid"
    h3 = ["Comparison", "U", "p (adjusted)", "r"]
    for i, v in enumerate(h3):
        t3.rows[0].cells[i].text = v
    rows3 = [
        ["Website vs YouTube", "1646", ".038", "-0.316"],
        ["Website vs TikTok", "2088", "< .001", "-0.670"],
        ["Website vs Instagram", "2364", "< .001", "-0.891"],
        ["YouTube vs TikTok", "1859", "< .001", "-0.487"],
        ["YouTube vs Instagram", "2313", "< .001", "-0.850"],
        ["TikTok vs Instagram", "1883", "< .001", "-0.506"],
    ]
    for ri, rd in enumerate(rows3):
        for ci, v in enumerate(rd):
            t3.rows[ri+1].cells[ci].text = v
    format_table(t3)
    shade_header_row(t3)

    add_body(doc, "")

    # -- Table C4: JAMA --
    add_heading2(doc, "Table C4. JAMA Benchmark Compliance by Platform (%)")
    t4 = doc.add_table(rows=5, cols=5)
    t4.style = "Table Grid"
    h4 = ["Criterion", "Website", "YouTube", "TikTok", "Instagram"]
    for i, v in enumerate(h4):
        t4.rows[0].cells[i].text = v
    rows4 = [
        ["Authorship", "30%", "38%", "30%", "16%"],
        ["Attribution", "30%", "16%", "0%", "4%"],
        ["Disclosure", "16%", "0%", "14%", "0%"],
        ["Currency", "68%", "100%", "100%", "100%"],
    ]
    for ri, rd in enumerate(rows4):
        for ci, v in enumerate(rd):
            t4.rows[ri+1].cells[ci].text = v
    format_table(t4)
    shade_header_row(t4)

    add_body(doc, "")

    # -- Table C5: Inter-rater reliability --
    add_heading2(doc, "Table C5. Inter-Rater Reliability Statistics")
    t5 = doc.add_table(rows=6, cols=3)
    t5.style = "Table Grid"
    h5 = ["Measure", "Statistic", "Interpretation"]
    for i, v in enumerate(h5):
        t5.rows[0].cells[i].text = v
    rows5 = [
        ["DISCERN Total (ICC[3,1])", "0.932 (95% CI [0.870, 0.960])", "Excellent"],
        ["JAMA Authorship (kappa)", "0.671", "Substantial"],
        ["JAMA Attribution (kappa)", "0.615", "Substantial"],
        ["JAMA Disclosure (kappa)", "0.273", "Fair (kappa paradox)"],
        ["JAMA Currency (kappa)", "0.216", "Fair (kappa paradox)"],
    ]
    for ri, rd in enumerate(rows5):
        for ci, v in enumerate(rd):
            t5.rows[ri+1].cells[ci].text = v
    format_table(t5)
    shade_header_row(t5)

    add_body(doc, "")

    # -- Table C6: Engagement correlations --
    add_heading2(doc, "Table C6. Spearman Correlations Between Engagement Metrics and DISCERN Scores")
    t6 = doc.add_table(rows=10, cols=4)
    t6.style = "Table Grid"
    h6 = ["Metric / Platform", "rho", "p", "Interpretation"]
    for i, v in enumerate(h6):
        t6.rows[0].cells[i].text = v
    rows6 = [
        ["Views (overall)", "-0.073", ".571", "Negligible, ns"],
        ["Likes (overall)", "0.413", "< .001", "Moderate positive (confounded)"],
        ["Comments (overall)", "0.415", "< .001", "Moderate positive (confounded)"],
        ["YouTube views", "-0.316", ".035", "Negative, significant"],
        ["YouTube likes", "-0.279", ".063", "ns"],
        ["TikTok views", "0.249", ".318", "ns"],
        ["TikTok likes", "0.280", ".098", "ns"],
        ["Instagram likes", "-0.013", ".936", "ns"],
        ["Note", "", "", "Overall likes/comments driven by Simpson's paradox"],
    ]
    for ri, rd in enumerate(rows6):
        for ci, v in enumerate(rd):
            t6.rows[ri+1].cells[ci].text = v
    format_table(t6)
    shade_header_row(t6)

    add_body(doc, "")

    # -- Table C7: Section scores --
    add_heading2(doc, "Table C7. DISCERN Section Scores by Platform (Median, IQR)")
    t7 = doc.add_table(rows=5, cols=5)
    t7.style = "Table Grid"
    h7 = ["Platform", "Section 1 Mdn", "Section 1 IQR", "Section 2 Mdn", "Section 2 IQR"]
    for i, v in enumerate(h7):
        t7.rows[0].cells[i].text = v
    rows7 = [
        ["Website", "25", "23\u201329", "18", "16\u201321"],
        ["YouTube", "24", "21\u201327", "16", "14\u201318"],
        ["TikTok", "19", "17\u201322", "14", "11\u201316"],
        ["Instagram", "15", "13\u201318", "10", "9\u201313"],
    ]
    for ri, rd in enumerate(rows7):
        for ci, v in enumerate(rd):
            t7.rows[ri+1].cells[ci].text = v
    format_table(t7)
    shade_header_row(t7)

    add_body(doc, "")

    # -- Table C8: Creator type --
    add_heading2(doc, "Table C8. DISCERN Scores by Creator Type")
    t8 = doc.add_table(rows=6, cols=3)
    t8.style = "Table Grid"
    h8 = ["Creator Type", "n", "Mdn DISCERN"]
    for i, v in enumerate(h8):
        t8.rows[0].cells[i].text = v
    rows8 = [
        ["Healthcare Professional", "17", "41.0"],
        ["Certified Fitness Professional", "31", "41.0"],
        ["Fitness Influencer", "27", "36.0"],
        ["General User", "30", "33.0"],
        ["Organisation", "85", "39.0"],
    ]
    for ri, rd in enumerate(rows8):
        for ci, v in enumerate(rd):
            t8.rows[ri+1].cells[ci].text = v
    format_table(t8)
    shade_header_row(t8)

    add_body(doc, "")

    # -- Table C9: Content format --
    add_heading2(doc, "Table C9. DISCERN Scores by Content Format")
    t9 = doc.add_table(rows=4, cols=4)
    t9.style = "Table Grid"
    h9 = ["Format", "n", "Mdn", "M"]
    for i, v in enumerate(h9):
        t9.rows[0].cells[i].text = v
    rows9 = [
        ["Long Video (>5 min)", "45", "43.0", "43.1"],
        ["Short Video (<60 s)", "18", "31.5", "32.8"],
        ["Carousel", "50", "27.5", "29.0"],
    ]
    for ri, rd in enumerate(rows9):
        for ci, v in enumerate(rd):
            t9.rows[ri+1].cells[ci].text = v
    format_table(t9)
    shade_header_row(t9)

    doc.save(os.path.join(APPENDICES, "appendix_c_data_summary.docx"))
    print("Appendix C saved.")


# ============================================================================
#  MAIN
# ============================================================================

if __name__ == "__main__":
    os.makedirs(CHAPTERS, exist_ok=True)
    os.makedirs(APPENDICES, exist_ok=True)

    print("Generating Chapter 3 (Research Findings)...")
    build_chapter3()

    print("Generating Chapter 4 (Discussion)...")
    build_chapter4()

    print("Generating Chapter 5 (Conclusions and Recommendations)...")
    build_chapter5()

    print("Generating Appendix C (Data Summary Tables)...")
    build_appendix_c()

    print("\nAll documents generated successfully.")
    print(f"  - {os.path.join(CHAPTERS, 'chapter3_findings.docx')}")
    print(f"  - {os.path.join(CHAPTERS, 'chapter4_discussion.docx')}")
    print(f"  - {os.path.join(CHAPTERS, 'chapter5_conclusions.docx')}")
    print(f"  - {os.path.join(APPENDICES, 'appendix_c_data_summary.docx')}")
