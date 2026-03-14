#!/usr/bin/env python3
"""
Generate Chapter 2: Research Methodology for Master's Thesis
Lithuanian Sports University
Author: Ayokunle Ademola-John
"""

import subprocess
import sys

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color):
    """Set background shading for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right, insideH, insideV."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val["val"]}" w:sz="{val["sz"]}" '
            f'w:space="0" w:color="{val["color"]}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_table_borders(table):
    """Apply borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def format_run(run, bold=False, italic=False, size=12, font_name="Times New Roman", color=None):
    """Format a run with specified properties."""
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_formatted_paragraph(doc, text, bold=False, italic=False, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=0, space_after=6, line_spacing=1.5,
                             first_line_indent=None):
    """Add a fully formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    format_run(run, bold=bold, italic=italic, size=size)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_heading_custom(doc, text, level=1, size=14, space_before=12, space_after=6):
    """Add a custom heading with Times New Roman formatting."""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    format_run(run, bold=True, size=size)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=1.5)
    return p


def add_apa_table_title(doc, table_number, title_text):
    """Add APA-style table title (number in bold, title in bold italic)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p.add_run(f"Table {table_number}")
    format_run(run1, bold=True, italic=False, size=12)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(title_text)
    format_run(run2, bold=False, italic=True, size=12)
    add_paragraph_spacing(p, before=12, after=0, line_spacing=1.5)
    add_paragraph_spacing(p2, before=0, after=6, line_spacing=1.5)
    return p, p2


def format_table_cell(cell, text, bold=False, italic=False, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Format a table cell with text."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    format_run(run, bold=bold, italic=italic, size=size)
    # Reduce paragraph spacing in cells
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15


def add_mixed_paragraph(doc, parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=0, space_after=6, line_spacing=1.5,
                         first_line_indent=None):
    """
    Add a paragraph with mixed formatting.
    parts: list of tuples (text, bold, italic)
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    for text, bold, italic in parts:
        run = p.add_run(text)
        format_run(run, bold=bold, italic=italic, size=12)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_bullet_point(doc, text, bold_prefix=None, indent_level=0):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        run_bold = p.add_run(bold_prefix)
        format_run(run_bold, bold=True, size=12)
        run_text = p.add_run(text)
        format_run(run_text, bold=False, size=12)
    else:
        run_text = p.add_run(text)
        format_run(run_text, bold=False, size=12)
    add_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)
    # Set bullet indentation
    pPr = p._element.get_or_add_pPr()
    indent = parse_xml(
        f'<w:ind {nsdecls("w")} w:left="{720 + indent_level * 360}" w:hanging="360"/>'
    )
    pPr.append(indent)
    # Add bullet character
    numPr = parse_xml(
        f'<w:numPr {nsdecls("w")}>'
        f'  <w:ilvl w:val="{indent_level}"/>'
        f'  <w:numId w:val="0"/>'
        f'</w:numPr>'
    )
    # Use a simple dash as bullet
    p.text = ""
    if bold_prefix:
        dash_run = p.add_run("\u2022 ")
        format_run(dash_run, bold=False, size=12)
        run_bold = p.add_run(bold_prefix)
        format_run(run_bold, bold=True, size=12)
        run_text = p.add_run(text)
        format_run(run_text, bold=False, size=12)
    else:
        dash_run = p.add_run("\u2022 ")
        format_run(dash_run, bold=False, size=12)
        run_text = p.add_run(text)
        format_run(run_text, bold=False, size=12)
    return p


def add_simple_bullet(doc, text, indent_cm=1.27):
    """Add a simple bullet point with dash."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    format_run(run, bold=False, size=12)
    add_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.first_line_indent = Cm(-0.63)
    return p


# ============================================================
# MAIN DOCUMENT GENERATION
# ============================================================

def create_chapter2():
    doc = Document()

    # ---- Page Setup ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)  # A4
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # Ensure Times New Roman for all character sets
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>')
    rPr.insert(0, rFonts)

    # ================================================================
    # CHAPTER 2 MAIN HEADING
    # ================================================================
    add_heading_custom(doc, "CHAPTER 2. RESEARCH METHODOLOGY", level=1, size=14, space_before=0, space_after=12)

    # Introductory paragraph
    add_formatted_paragraph(
        doc,
        "This chapter presents the methodological framework employed to investigate the quality of physical activity (PA) information across traditional websites and social media platforms. The research methodology encompasses the research object, strategy and logic, nature of the research, description of research subjects, data collection and assessment methods, organisational procedures, and statistical analysis techniques. Each component is detailed to ensure transparency, replicability, and methodological rigour in alignment with established standards for content analysis in health communication research.",
        first_line_indent=1.27
    )

    # ================================================================
    # 2.1 RESEARCH OBJECT
    # ================================================================
    add_heading_custom(doc, "2.1 Research Object", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "The research object of the present study is the quality of physical activity information available on traditional websites and social media platforms. Specifically, this investigation examines health-related content pertaining to physical activity that is publicly accessible through Google search results (representing traditional web-based information), YouTube, TikTok, and Instagram (representing contemporary social media platforms). Physical activity information is defined as any content that provides guidance, recommendations, instructions, or educational material regarding exercise, fitness, or physical activity behaviours intended for a general adult audience.",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "The assessment of information quality is operationalised through two established benchmarks: the DISCERN instrument (Charnock et al., 1999), which evaluates the reliability and completeness of health information, and the JAMA benchmarks (Silberg et al., 1997), which assess accountability standards including authorship, attribution, disclosure, and currency. Together, these instruments provide a comprehensive evaluation framework that captures both the substantive quality and the transparency of health information across diverse digital platforms.",
        first_line_indent=1.27
    )

    # ================================================================
    # 2.2 RESEARCH STRATEGY AND LOGIC
    # ================================================================
    add_heading_custom(doc, "2.2 Research Strategy and Logic", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "The present study employs a cross-sectional content analysis design, which enables the systematic examination and comparison of physical activity information quality across four digital platforms at a single point in time. Content analysis is a well-established research method in health communication that facilitates the objective, systematic, and quantitative description of manifest content (Krippendorff, 2018). The cross-sectional approach was selected for its capacity to capture a snapshot of the current information landscape, allowing for direct comparisons across platforms under identical assessment conditions.",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "The research follows a deductive approach, applying established quality assessment frameworks\u2014the DISCERN instrument and JAMA benchmarks\u2014to evaluate content across a novel four-platform comparison. While these instruments were originally developed for traditional health information materials, their application to social media content represents a theoretically grounded extension that has been validated in prior research examining YouTube (Azer, 2020), TikTok (Kong et al., 2021), and Instagram health content (Chea & Lim, 2023).",
        first_line_indent=1.27
    )

    # SPIDER framework
    add_formatted_paragraph(
        doc,
        "The study design was structured using the SPIDER framework (Cooke et al., 2012), adapted for content analysis as follows:",
        first_line_indent=1.27
    )

    spider_items = [
        ("\u2022 Sample: ", "200 content items across four platforms (50 per platform)."),
        ("\u2022 Phenomenon of Interest: ", "Quality of physical activity information."),
        ("\u2022 Design: ", "Cross-sectional content analysis."),
        ("\u2022 Evaluation: ", "DISCERN instrument and JAMA benchmarks."),
        ("\u2022 Research type: ", "Quantitative."),
    ]
    for bold_part, normal_part in spider_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run_b = p.add_run(bold_part)
        format_run(run_b, bold=True, italic=False, size=12)
        run_n = p.add_run(normal_part)
        format_run(run_n, bold=False, italic=False, size=12)
        add_paragraph_spacing(p, before=0, after=2, line_spacing=1.5)
        p.paragraph_format.left_indent = Cm(1.27)

    add_formatted_paragraph(
        doc,
        "The rationale for employing a cross-sectional content analysis rather than a systematic review lies in the current state of the literature. Recent meta-analyses have already synthesised the existing body of evidence on health information quality across individual platforms (Ayyash et al., 2025; Ayyash, Musaad, et al., 2025). However, no study to date has conducted a direct, head-to-head comparison of physical activity information quality across traditional websites, YouTube, TikTok, and Instagram using both the DISCERN and JAMA instruments simultaneously. The present study addresses this gap by providing the first four-platform comparative analysis using dual quality assessment frameworks applied specifically to physical activity content.",
        first_line_indent=1.27
    )

    # ================================================================
    # 2.3 NATURE OF RESEARCH
    # ================================================================
    add_heading_custom(doc, "2.3 Nature of Research", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "The present study is quantitative in nature, employing numerical scoring instruments to assess information quality and statistical methods to analyse the resulting data. The research is both descriptive and comparative: descriptive in that it characterises the quality of physical activity information across platforms using standardised metrics, and comparative in that it systematically examines differences in quality scores between platform types, creator categories, and content formats.",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "Furthermore, this study represents applied research with practical implications for multiple stakeholders. The findings are intended to inform health literacy initiatives, guide content regulation policies on social media platforms, assist healthcare professionals in directing patients toward reliable information sources, and contribute to the broader understanding of digital health information ecosystems. The applied orientation aligns with the growing recognition that the quality of online health information has direct consequences for public health behaviours and outcomes (Swire-Thompson & Lazer, 2020).",
        first_line_indent=1.27
    )

    # ================================================================
    # 2.4 CONTINGENT OF RESEARCH SUBJECTS
    # ================================================================
    add_heading_custom(doc, "2.4 Study Sample and Sampling Frame", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "The study sample comprises a total of N = 200 content items collected from four digital platforms: traditional websites accessed through Google search (n = 50), YouTube (n = 50), TikTok (n = 50), and Instagram (n = 50). This sample size was determined based on precedent in health information quality research, where samples of 40\u201360 items per platform have been established as sufficient for detecting meaningful quality differences (Leong et al., 2021; Drozd et al., 2018).",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "Content was retrieved using five standardised search terms designed to simulate typical user queries related to physical activity:",
        first_line_indent=1.27
    )

    search_terms = [
        "\u2022 \u201cHow to start exercising\u201d",
        "\u2022 \u201cBest exercises to lose weight\u201d",
        "\u2022 \u201cStrength training for beginners\u201d",
        "\u2022 \u201cPhysical activity guidelines\u201d",
        "\u2022 \u201cHome workout routine\u201d"
    ]
    for term in search_terms:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(term)
        format_run(run, size=12)
        add_paragraph_spacing(p, before=0, after=2, line_spacing=1.5)
        p.paragraph_format.left_indent = Cm(1.27)

    add_formatted_paragraph(
        doc,
        "These search terms were selected to represent a range of common physical activity information-seeking behaviours, from general exercise initiation to specific activity modalities. For each search term, the first 10 results were collected from each platform, yielding 10 results \u00d7 5 search terms \u00d7 4 platforms = 200 content items. The sampling strategy employed convenience sampling of top-ranked search results, which simulates the actual information-seeking behaviour of typical users who predominantly access the first page of search results (Pan et al., 2007).",
        first_line_indent=1.27
    )

    # Inclusion / Exclusion criteria
    add_mixed_paragraph(
        doc,
        [("Inclusion criteria. ", True, False),
         ("Content items were included if they met all of the following criteria: (a) written or presented in English, (b) publicly accessible without account registration or payment, (c) containing substantive physical activity information (i.e., providing guidance, instructions, or educational content about exercise or physical activity), and (d) posted or last updated within the three years preceding data collection (i.e., from February 2023 onward).", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Exclusion criteria. ", True, False),
         ("Content items were excluded if they met any of the following criteria: (a) duplicate content appearing across multiple platforms, (b) content behind a paywall or requiring subscription access, (c) content not primarily focused on physical activity (e.g., product-only advertisements, unrelated entertainment), or (d) content in languages other than English. When a result was excluded, the next sequential result from the same search was included as a replacement to maintain the target sample size.", False, False)],
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "Table 2.1 presents the sampling matrix illustrating the distribution of content items across search terms and platforms.",
        first_line_indent=1.27
    )

    # TABLE 2.1 - Sampling Matrix
    add_apa_table_title(doc, "2.1", "Sampling Matrix: Distribution of Content Items by Search Term and Platform")

    table1 = doc.add_table(rows=7, cols=6)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table1)

    # Header row
    headers1 = ["Search Term", "Google\n(n)", "YouTube\n(n)", "TikTok\n(n)", "Instagram\n(n)", "Total\n(N)"]
    for i, header in enumerate(headers1):
        format_table_cell(table1.rows[0].cells[i], header, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table1.rows[0].cells[i], "D9E2F3")

    # Data rows
    terms_short = [
        "How to start exercising",
        "Best exercises to lose weight",
        "Strength training for beginners",
        "Physical activity guidelines",
        "Home workout routine"
    ]
    for row_idx, term in enumerate(terms_short, start=1):
        format_table_cell(table1.rows[row_idx].cells[0], term, size=10)
        for col in range(1, 5):
            format_table_cell(table1.rows[row_idx].cells[col], "10", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_table_cell(table1.rows[row_idx].cells[5], "40", size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Total row
    format_table_cell(table1.rows[6].cells[0], "Total", bold=True, size=10)
    for col in range(1, 5):
        format_table_cell(table1.rows[6].cells[col], "50", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    format_table_cell(table1.rows[6].cells[5], "200", bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(6):
        set_cell_shading(table1.rows[6].cells[i], "E2EFDA")

    # Set column widths for table 2.1
    for row in table1.rows:
        row.cells[0].width = Cm(5.5)
        for col_idx in range(1, 6):
            row.cells[col_idx].width = Cm(2.2)

    # Note below table
    p_note = doc.add_paragraph()
    run_note = p_note.add_run("Note. ")
    format_run(run_note, bold=False, italic=True, size=10)
    run_note2 = p_note.add_run("Each cell represents the number of content items collected per search term per platform. Total sample N = 200.")
    format_run(run_note2, bold=False, italic=False, size=10)
    add_paragraph_spacing(p_note, before=3, after=12, line_spacing=1.0)

    # ================================================================
    # 2.5 RESEARCH METHODS
    # ================================================================
    add_heading_custom(doc, "2.5 Research Methods", level=2, size=12, space_before=12, space_after=6)

    # 2.5.1 Data Collection
    add_heading_custom(doc, "2.5.1 Data Collection Procedures", level=3, size=12, space_before=8, space_after=6)

    add_formatted_paragraph(
        doc,
        "Data collection was conducted over a two-day window (1\u20132 March 2026) using a systematic protocol designed to minimise algorithmic bias and ensure reproducibility. All searches were performed in logged-out, incognito (private) browsing mode in Google Chrome to eliminate the influence of personalised search algorithms, browsing history, and account-based recommendations on result rankings (Tripodi, 2018). The browser\u2019s language was set to English (United States) and geolocation was configured to Kaunas, Lithuania, the researcher\u2019s institutional location, to standardise the search environment across all platforms. No VPN was used.",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "For traditional websites, searches were conducted on Google (google.com) using the default \u2018All\u2019 results tab; sponsored results and advertisements were excluded. For YouTube, searches were conducted on youtube.com using the default \u2018Relevance\u2019 sort. For both platforms, an automated data collection approach was employed using Playwright, an open-source browser automation framework. The automated scripts navigated to each platform, entered the standardised search terms, and systematically recorded the top 10 organic results for each query. This approach ensured consistency in result capture and eliminated human error in the recording of search rankings, URLs, and metadata. For TikTok and Instagram, manual data collection was conducted due to these platforms\u2019 implementation of anti-bot measures that prevent reliable browser automation. TikTok searches used the in-app \u2018Top\u2019 tab; Instagram searches used the \u2018Tags\u2019 and \u2018Top\u2019 tabs. The primary researcher manually searched each term within the platform\u2019s native search function, recording the first 10 relevant results in order of appearance. Reposts, advertisements, pinned/promoted content, and duplicate items from the same creator for the same search term were excluded and replaced by the next sequential organic result.",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "All search results were recorded in their order of appearance within the search ranking. For each content item, the following information was documented: platform, search term used, ranking position, URL or content identifier, creator name, creator type classification, content format, date of publication or last update, and engagement metrics (views, likes, shares, and comments where available). Screenshots were captured for each content item to establish an audit trail, supporting the verification and reliability of the assessment process.",
        first_line_indent=1.27
    )

    # 2.5.2 Quality Assessment Instruments
    add_heading_custom(doc, "2.5.2 Quality Assessment Instruments", level=3, size=12, space_before=8, space_after=6)

    # DISCERN section
    add_mixed_paragraph(
        doc,
        [("DISCERN Instrument. ", True, False),
         ("The DISCERN instrument (Charnock et al., 1999) is a validated and widely used tool for assessing the quality of written consumer health information. The instrument comprises 16 items, each scored on a Likert-type scale ranging from 1 (definite no) to 5 (definite yes). The instrument is divided into three sections: Section 1 (Questions 1\u20138) evaluates the reliability of the publication, addressing whether the content has clear aims, achieves those aims, is relevant, identifies its sources, is balanced, and provides sufficient detail for decision-making. Section 2 (Questions 9\u201315) assesses the quality of information about treatment choices, including descriptions of how treatments work, their benefits, risks, effects on quality of life, and what would happen without treatment. Section 3 comprises a single item (Question 16) providing an overall quality rating. Total scores range from 16 to 80.", False, False)],
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "For the purposes of the present study, DISCERN quality scores were categorised using the following established thresholds: Very Poor (16\u201326), Poor (27\u201338), Fair (39\u201350), Good (51\u201362), and Excellent (63\u201380). These categories facilitate meaningful interpretation and cross-study comparison of quality ratings. In adapting the instrument for social media content, questions pertaining to treatment information were interpreted in the context of exercise and physical activity recommendations, consistent with the approach adopted in prior studies examining fitness-related content (Kocyigit et al., 2019; Kong et al., 2021).",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "Table 2.2 presents the complete DISCERN instrument with all 16 assessment questions.",
        first_line_indent=1.27
    )

    # TABLE 2.2 - DISCERN Questions
    add_apa_table_title(doc, "2.2", "DISCERN Instrument Assessment Questions")

    table2 = doc.add_table(rows=18, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table2)

    # Header
    t2_headers = ["Item", "Section", "Question"]
    for i, h in enumerate(t2_headers):
        format_table_cell(table2.rows[0].cells[i], h, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table2.rows[0].cells[i], "D9E2F3")

    discern_questions = [
        ("Q1", "1: Reliability", "Are the aims clear?"),
        ("Q2", "1: Reliability", "Does it achieve its aims?"),
        ("Q3", "1: Reliability", "Is it relevant?"),
        ("Q4", "1: Reliability", "Is it clear what sources of information were used to compile the publication?"),
        ("Q5", "1: Reliability", "Is it clear when the information used or reported in the publication was produced?"),
        ("Q6", "1: Reliability", "Is it balanced and unbiased?"),
        ("Q7", "1: Reliability", "Does it provide details of additional sources of support and information?"),
        ("Q8", "1: Reliability", "Does it refer to areas of uncertainty?"),
        ("Q9", "2: Treatment Quality", "Does it describe how each treatment works?"),
        ("Q10", "2: Treatment Quality", "Does it describe the benefits of each treatment?"),
        ("Q11", "2: Treatment Quality", "Does it describe the risks of each treatment?"),
        ("Q12", "2: Treatment Quality", "Does it describe what would happen if no treatment is used?"),
        ("Q13", "2: Treatment Quality", "Does it describe the effects of treatment choices on overall quality of life?"),
        ("Q14", "2: Treatment Quality", "Is it clear that there may be more than one possible treatment choice?"),
        ("Q15", "2: Treatment Quality", "Does it provide support for shared decision-making?"),
        ("Q16", "3: Overall Rating", "Based on the answers to all of the above questions, rate the overall quality of the publication."),
    ]
    # Section heading row
    row_start = 1
    for idx, (item, section, question) in enumerate(discern_questions):
        row = table2.rows[row_start + idx]
        format_table_cell(row.cells[0], item, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_table_cell(row.cells[1], section, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_table_cell(row.cells[2], question, size=10)

    # Set column widths
    for row in table2.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(11.0)

    # Table note
    p_note2 = doc.add_paragraph()
    run_n2a = p_note2.add_run("Note. ")
    format_run(run_n2a, italic=True, size=10)
    run_n2b = p_note2.add_run("Adapted from \u201cDISCERN: An instrument for judging the quality of written consumer health information on treatment choices,\u201d by D. Charnock, S. Shepperd, G. Needham, and R. Gann, 1999, Journal of Epidemiology & Community Health, 53(2), 105\u2013111. Each item is scored from 1 (definite no) to 5 (definite yes). For social media content, treatment-related questions (Q9\u2013Q15) were interpreted as referring to exercise and physical activity recommendations.")
    format_run(run_n2b, size=10)
    add_paragraph_spacing(p_note2, before=3, after=12, line_spacing=1.0)

    # JAMA Benchmarks
    add_mixed_paragraph(
        doc,
        [("JAMA Benchmarks. ", True, False),
         ("The JAMA benchmarks (Silberg et al., 1997) provide a concise accountability framework for evaluating health information. The instrument assesses four criteria, each scored dichotomously as present (Y) or absent (N): (a) Authorship\u2014whether authors and their credentials or affiliations are clearly identified; (b) Attribution\u2014whether references and sources for content are cited; (c) Disclosure\u2014whether sponsorship, advertising, conflicts of interest, or financial relationships are declared; and (d) Currency\u2014whether the date of content creation or most recent update is provided. The total JAMA score ranges from 0 to 4, representing the number of criteria met.", False, False)],
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "To apply the JAMA benchmarks to social media content, which differs substantially in format from traditional web-based health information, the criteria were operationalised as described in Table 2.3. This operationalisation draws on adaptations employed in prior research assessing social media health content (Delli et al., 2016; Kocyigit et al., 2019).",
        first_line_indent=1.27
    )

    # TABLE 2.3 - JAMA Benchmarks Operationalisation
    add_apa_table_title(doc, "2.3", "JAMA Benchmarks Criteria and Social Media Operationalisation")

    table3 = doc.add_table(rows=5, cols=3)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table3)

    t3_headers = ["Criterion", "Original Definition", "Social Media Operationalisation"]
    for i, h in enumerate(t3_headers):
        format_table_cell(table3.rows[0].cells[i], h, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table3.rows[0].cells[i], "D9E2F3")

    jama_data = [
        ("Authorship", "Authors and credentials/affiliations identified",
         "Creator bio states professional credentials (e.g., MD, PT, CSCS) or account is a verified professional/organisational account"),
        ("Attribution", "References and sources for content are cited",
         "Content cites scientific studies, clinical guidelines, or other sources in video, caption, description, or on-screen text references"),
        ("Disclosure", "Sponsorship, advertising, or conflicts of interest declared",
         "Content discloses sponsorships, partnerships, or paid promotions (e.g., #ad, #sponsored, platform-provided paid partnership label)"),
        ("Currency", "Date of content creation or update provided",
         "Post or publication date is visible (inherent feature of most social media platforms; for websites, date must be explicitly displayed)")
    ]
    for idx, (crit, orig, sm) in enumerate(jama_data):
        row = table3.rows[idx + 1]
        format_table_cell(row.cells[0], crit, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        format_table_cell(row.cells[1], orig, size=10)
        format_table_cell(row.cells[2], sm, size=10)

    for row in table3.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(5.5)
        row.cells[2].width = Cm(8.0)

    p_note3 = doc.add_paragraph()
    run_n3a = p_note3.add_run("Note. ")
    format_run(run_n3a, italic=True, size=10)
    run_n3b = p_note3.add_run("Adapted from \u201cAssessing, controlling, and assuring the quality of medical information on the internet,\u201d by W. M. Silberg, G. D. Lundberg, and R. A. Musacchio, 1997, JAMA, 277(15), 1244\u20131245. Each criterion is scored as present (Y) or absent (N). CSCS = Certified Strength and Conditioning Specialist; MD = Medical Doctor; PT = Physical Therapist.")
    format_run(run_n3b, size=10)
    add_paragraph_spacing(p_note3, before=3, after=12, line_spacing=1.0)

    # Additional Variables
    add_heading_custom(doc, "2.5.3 Additional Variables", level=3, size=12, space_before=8, space_after=6)

    add_formatted_paragraph(
        doc,
        "In addition to the primary quality assessment instruments, several supplementary variables were recorded for each content item to facilitate secondary analyses examining the moderating role of creator characteristics and content features on information quality.",
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Creator type classification. ", True, False),
         ("Each content creator was classified into one of five categories based on the information available in their profile, bio, and content: (a) Healthcare professional (e.g., physicians, physiotherapists, registered nurses), (b) Certified fitness professional (e.g., certified personal trainers, strength and conditioning specialists), (c) Fitness influencer (individuals with substantial followings who create fitness content without verified professional credentials), (d) General user (individuals without identifiable fitness or health credentials), and (e) Organisation (institutional accounts such as hospitals, health agencies, or fitness organisations). This classification scheme was developed based on typologies employed in prior research on health information creators (Kocyigit et al., 2019; Yurdaisik, 2022).", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Engagement metrics. ", True, False),
         ("Where available, the following engagement metrics were recorded at the time of data collection: total views, likes (or equivalent reactions), shares, and comments. These metrics serve as indicators of content reach and audience interaction, enabling the examination of the relationship between content popularity and information quality (Osman et al., 2022).", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Content format. ", True, False),
         ("Each content item was classified by format: (a) Text article (traditional websites), (b) Long-form video (>5 minutes; primarily YouTube), (c) Short-form video (<60 seconds; primarily TikTok and Instagram Reels), (d) Image with caption (primarily Instagram posts), and (e) Carousel (multi-image Instagram posts). This classification enables the examination of whether content format is associated with information quality.", False, False)],
        first_line_indent=1.27
    )

    # ================================================================
    # 2.6 RESEARCH ORGANISATION
    # ================================================================
    add_heading_custom(doc, "2.6 Research Organisation", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "The study was conducted over a defined period from February to March 2026, with data collection completed within a concentrated two-day window to minimise temporal variation in search results and engagement metrics. The research process was organised into six sequential phases, as illustrated in Figure 2.1.",
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Phase 1: Automated data collection (Day 1). ", True, False),
         ("Playwright browser automation scripts were executed to collect search results from Google and YouTube. The scripts performed searches for all five search terms, captured the top 10 results per term, and recorded URLs, metadata, and screenshots.", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Phase 2: Manual data collection (Days 1\u20132). ", True, False),
         ("The primary researcher manually collected search results from TikTok and Instagram using the platform-native search functions. Results were recorded in the same standardised format as the automated collection.", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Phase 3: Primary quality assessment. ", True, False),
         ("The primary researcher assessed all 200 content items using both the DISCERN instrument and the JAMA benchmarks. For text-based website content, each article was read in full. For video content (YouTube and TikTok), scoring was based on a combination of the video content itself, visible captions, descriptions, and on-screen text; full verbatim transcription was not performed. For Instagram carousels and image posts, scoring was based on all visible text, captions, and associated profile information. This approach reflects the information a typical user would encounter when engaging with each item. Scores were recorded in a structured data collection spreadsheet.", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Phase 4: Second rater assessment. ", True, False),
         ("A trained second rater independently scored a 20% subsample of content items (n = 40; 10 randomly selected items per platform) using both instruments. The second rater was provided with the same scoring guidelines and training materials as the primary researcher to ensure standardised application of the assessment criteria.", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Phase 5: Inter-rater reliability assessment. ", True, False),
         ("Agreement between the primary researcher and the second rater was evaluated using the Intraclass Correlation Coefficient (ICC) for DISCERN scores and Cohen\u2019s kappa (\u03ba) for JAMA criteria. These metrics are reported in the Results chapter.", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Phase 6: Statistical analysis. ", True, False),
         ("Following the completion of quality assessments, all data were compiled and analysed using Python 3.13 with scientific computing libraries, as detailed in Section 2.7.", False, False)],
        first_line_indent=1.27
    )

    # Figure 2.1 placeholder
    add_formatted_paragraph(
        doc,
        "",
        space_before=6, space_after=0
    )

    # Create a simple flow diagram representation as a table
    p_fig_title = doc.add_paragraph()
    p_fig_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fig = p_fig_title.add_run("Figure 2.1")
    format_run(run_fig, bold=True, italic=False, size=12)
    add_paragraph_spacing(p_fig_title, before=12, after=0, line_spacing=1.5)

    p_fig_desc = doc.add_paragraph()
    p_fig_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fig_d = p_fig_desc.add_run("Flow Diagram of the Research Process")
    format_run(run_fig_d, bold=False, italic=True, size=12)
    add_paragraph_spacing(p_fig_desc, before=0, after=6, line_spacing=1.5)

    # Flow diagram as vertically stacked boxes
    phases = [
        "Phase 1: Automated Data Collection\n(Google & YouTube via Playwright \u2013 Day 1)",
        "Phase 2: Manual Data Collection\n(TikTok & Instagram \u2013 Days 1\u20132)",
        "Phase 3: Primary DISCERN & JAMA Assessment\n(N = 200 content items)",
        "Phase 4: Second Rater Assessment\n(20% subsample, n = 40)",
        "Phase 5: Inter-Rater Reliability\n(ICC for DISCERN; Cohen\u2019s \u03ba for JAMA)",
        "Phase 6: Statistical Analysis\n(Python 3.13 with SciPy, pingouin, scikit-posthocs)"
    ]

    flow_table = doc.add_table(rows=11, cols=1)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, phase_text in enumerate(phases):
        row_idx = i * 2
        cell = flow_table.rows[row_idx].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(phase_text)
        format_run(run, bold=False, size=10)
        # Add border to phase cells
        set_cell_border(cell,
            top={"val": "single", "sz": "8", "color": "000000"},
            bottom={"val": "single", "sz": "8", "color": "000000"},
            left={"val": "single", "sz": "8", "color": "000000"},
            right={"val": "single", "sz": "8", "color": "000000"})

        # Add arrow between phases (except after last)
        if i < len(phases) - 1:
            arrow_row = row_idx + 1
            arrow_cell = flow_table.rows[arrow_row].cells[0]
            arrow_cell.text = ""
            p_arrow = arrow_cell.paragraphs[0]
            p_arrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_arrow = p_arrow.add_run("\u2193")
            format_run(run_arrow, bold=True, size=14)
            # No borders for arrow rows
            set_cell_border(arrow_cell,
                top={"val": "none", "sz": "0", "color": "FFFFFF"},
                bottom={"val": "none", "sz": "0", "color": "FFFFFF"},
                left={"val": "none", "sz": "0", "color": "FFFFFF"},
                right={"val": "none", "sz": "0", "color": "FFFFFF"})

    for row in flow_table.rows:
        row.cells[0].width = Cm(10)

    p_fig_note = doc.add_paragraph()
    p_fig_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fn = p_fig_note.add_run("Note. ")
    format_run(run_fn, italic=True, size=10)
    run_fn2 = p_fig_note.add_run("The research process proceeded sequentially through six phases over the study period (February\u2013March 2026).")
    format_run(run_fn2, size=10)
    add_paragraph_spacing(p_fig_note, before=3, after=12, line_spacing=1.0)

    # Ethics
    add_heading_custom(doc, "2.6.1 Ethical Considerations", level=3, size=12, space_before=8, space_after=6)

    add_formatted_paragraph(
        doc,
        "The present study does not involve human subjects, as all data were derived from publicly available online content. Consequently, formal ethics committee approval was not required, consistent with standard practice for content analyses of publicly accessible materials (Snelson, 2016). No personal data of content creators were collected beyond publicly displayed profile information relevant to the creator type classification. All content was accessed through standard public-facing interfaces, and no private or restricted materials were accessed at any point during the data collection process.",
        first_line_indent=1.27
    )

    # ================================================================
    # 2.7 METHODS OF STATISTICAL ANALYSIS
    # ================================================================
    add_heading_custom(doc, "2.7 Methods of Statistical Analysis", level=2, size=12, space_before=12, space_after=6)

    add_formatted_paragraph(
        doc,
        "All statistical analyses were conducted using Python 3.13 with the following scientific computing libraries: pandas (McKinney, 2010) for data manipulation, SciPy (Virtanen et al., 2020) for non-parametric inferential tests including Kruskal\u2013Wallis, Mann\u2013Whitney U, chi-square, and Spearman correlations, scikit-posthocs (Terpilowski, 2019) for Dunn\u2019s post-hoc comparisons with Bonferroni correction, pingouin (Vallat, 2018) for intraclass correlation coefficient computation, and matplotlib/seaborn for publication-quality graphical representations. The significance level was set at \u03b1 = 0.05 for all inferential tests. The complete analysis code is reproduced in Appendix D.",
        first_line_indent=1.27
    )

    # Descriptive statistics
    add_heading_custom(doc, "2.7.1 Descriptive Statistics", level=3, size=12, space_before=8, space_after=6)

    add_formatted_paragraph(
        doc,
        "Descriptive statistics were employed to characterise the distribution of quality scores across platforms and content categories. For DISCERN total scores, the median, interquartile range (IQR), and range were reported for each platform, as ordinal Likert-type data are best represented by median-based measures of central tendency (Sullivan & Artino, 2013). For JAMA benchmarks, frequencies and percentages of compliance for each criterion were calculated per platform, along with overall JAMA score distributions.",
        first_line_indent=1.27
    )

    # Inter-rater reliability
    add_heading_custom(doc, "2.7.2 Inter-Rater Reliability", level=3, size=12, space_before=8, space_after=6)

    add_formatted_paragraph(
        doc,
        "Inter-rater reliability was assessed for the 20% subsample (n = 40) independently scored by both the primary researcher and the second rater. For DISCERN total scores, the Intraclass Correlation Coefficient (ICC) was calculated using a two-way mixed-effects model, single measures, with absolute agreement (ICC[3,1]; Koo & Li, 2016). ICC values were interpreted according to established guidelines: values less than 0.50 indicate poor reliability, 0.50\u20130.75 indicate moderate reliability, 0.75\u20130.90 indicate good reliability, and values greater than 0.90 indicate excellent reliability (Koo & Li, 2016).",
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "For the JAMA benchmark criteria, which are scored dichotomously (present/absent), Cohen\u2019s kappa (\u03ba) was calculated for each criterion. Kappa values were interpreted using the guidelines proposed by Landis and Koch (1977): values below 0.20 indicate slight agreement, 0.21\u20130.40 indicate fair agreement, 0.41\u20130.60 indicate moderate agreement, 0.61\u20130.80 indicate substantial agreement, and values above 0.80 indicate almost perfect agreement.",
        first_line_indent=1.27
    )

    # Inferential statistics
    add_heading_custom(doc, "2.7.3 Inferential Statistics", level=3, size=12, space_before=8, space_after=6)

    add_formatted_paragraph(
        doc,
        "Non-parametric statistical tests were selected as the primary inferential methods due to the ordinal nature of DISCERN scores and the anticipated non-normal distribution of quality data, consistent with methodological recommendations for health information quality research (Naaman et al., 2022).",
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Platform comparisons of DISCERN scores. ", True, False),
         ("The Kruskal-Wallis H test was employed to examine whether DISCERN total scores differed significantly across the four platforms (Google, YouTube, TikTok, and Instagram). When a statistically significant omnibus result was obtained, Dunn\u2019s post-hoc test with Bonferroni correction was applied to identify which specific platform pairs differed significantly. The effect size was quantified using eta-squared (\u03b7\u00b2), computed as \u03b7\u00b2 = (H \u2212 k + 1) / (N \u2212 k), where H is the Kruskal-Wallis statistic, k the number of groups, and N the total sample size. This is an approximation rather than a true proportion of variance explained and should be interpreted accordingly; values of approximately 0.01, 0.06, and 0.14 correspond to small, medium, and large effects, respectively (Cohen, 1988).", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Platform comparisons of JAMA compliance. ", True, False),
         ("The chi-square test of independence was used to assess whether the proportion of content meeting each JAMA criterion differed across platforms. Where expected cell frequencies fell below 5, Fisher\u2019s exact test was employed as an alternative. The effect size was reported as Cram\u00e9r\u2019s V, interpreted according to Cohen\u2019s (1988) guidelines: values of 0.10, 0.30, and 0.50 represent small, medium, and large effects for tables with three degrees of freedom.", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Creator type and quality scores. ", True, False),
         ("To investigate whether creator type moderates information quality, DISCERN scores were compared between professional creators (healthcare professionals and certified fitness professionals combined) and non-professional creators (fitness influencers, general users, and organisations combined) using the Mann-Whitney U test. Organisations were classified as non-professional because they encompass a heterogeneous mix of entities\u2014from government health departments to commercial fitness brands\u2014making uniform professional classification inappropriate. The effect size was reported as rank-biserial correlation (r", False, False),
         ("rb", False, True),
         ("), where values of 0.10, 0.30, and 0.50 correspond to small, medium, and large effects (Fritz et al., 2012).", False, False)],
        first_line_indent=1.27
    )

    add_mixed_paragraph(
        doc,
        [("Engagement metrics and information quality. ", True, False),
         ("Spearman\u2019s rank correlation coefficient (\u03c1) was calculated to examine the relationship between engagement metrics (views, likes, shares) and DISCERN total scores. This non-parametric correlation was selected due to the highly skewed distribution typically observed in engagement data and the ordinal nature of DISCERN scores. Correlation strength was interpreted as: \u03c1 < 0.10 negligible, 0.10\u20130.39 weak, 0.40\u20130.69 moderate, and \u2265 0.70 strong (Schober et al., 2018).", False, False)],
        first_line_indent=1.27
    )

    add_formatted_paragraph(
        doc,
        "Table 2.4 provides a summary of the statistical analysis plan, linking each research question to the corresponding variables, statistical tests, and rationale for test selection.",
        first_line_indent=1.27
    )

    # TABLE 2.4 - Statistical Analysis Plan
    add_apa_table_title(doc, "2.4", "Summary of Statistical Analysis Plan by Research Question")

    table4 = doc.add_table(rows=5, cols=4)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table4)

    t4_headers = ["Research Question", "Variables", "Statistical Test", "Rationale"]
    for i, h in enumerate(t4_headers):
        format_table_cell(table4.rows[0].cells[i], h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table4.rows[0].cells[i], "D9E2F3")

    t4_data = [
        (
            "RQ1: How does PA information quality differ between traditional websites and social media?",
            "IV: Platform (4 levels)\nDV: DISCERN total score; JAMA criteria compliance",
            "Kruskal-Wallis H test (DISCERN)\nDunn\u2019s post-hoc with Bonferroni\n\u03c7\u00b2 / Fisher\u2019s exact test (JAMA)\nEffect sizes: \u03b7\u00b2, Cram\u00e9r\u2019s V",
            "Non-parametric; ordinal DISCERN data likely non-normally distributed. Chi-square appropriate for categorical JAMA data."
        ),
        (
            "RQ2: Does creator type moderate quality scores?",
            "IV: Creator type (professional vs. non-professional)\nDV: DISCERN total score",
            "Mann-Whitney U test\nEffect size: rank-biserial correlation",
            "Non-parametric comparison of two independent groups with ordinal outcome variable."
        ),
        (
            "RQ3: What is the relationship between engagement metrics and quality?",
            "Predictor: Views, likes, shares\nOutcome: DISCERN total score",
            "Spearman\u2019s rank correlation (\u03c1)",
            "Non-parametric correlation appropriate for skewed engagement data and ordinal quality scores."
        ),
        (
            "Inter-rater reliability",
            "Rater 1 vs. Rater 2 scores\n(20% subsample, n = 40)",
            "ICC (two-way mixed, single measures, absolute agreement) for DISCERN\nCohen\u2019s \u03ba for each JAMA criterion",
            "ICC suitable for continuous/ordinal data; \u03ba appropriate for dichotomous categorical data."
        ),
    ]

    for idx, (rq, var, test, rationale) in enumerate(t4_data):
        row = table4.rows[idx + 1]
        format_table_cell(row.cells[0], rq, size=9)
        format_table_cell(row.cells[1], var, size=9)
        format_table_cell(row.cells[2], test, size=9)
        format_table_cell(row.cells[3], rationale, size=9)

    for row in table4.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(4.5)
        row.cells[3].width = Cm(4.0)

    p_note4 = doc.add_paragraph()
    run_n4a = p_note4.add_run("Note. ")
    format_run(run_n4a, italic=True, size=10)
    run_n4b = p_note4.add_run("IV = independent variable; DV = dependent variable; PA = physical activity; ICC = Intraclass Correlation Coefficient; \u03b7\u00b2 = eta-squared (approximation); \u03c7\u00b2 = chi-square; \u03ba = kappa; \u03c1 = Spearman\u2019s rho. Significance level: \u03b1 = 0.05 for all tests.")
    format_run(run_n4b, size=10)
    add_paragraph_spacing(p_note4, before=3, after=12, line_spacing=1.0)

    # Concluding paragraph for the chapter
    add_formatted_paragraph(
        doc,
        "In summary, the methodological approach described in this chapter was designed to provide a rigorous, transparent, and replicable framework for evaluating the quality of physical activity information across four major digital platforms. The combination of the DISCERN instrument and JAMA benchmarks offers a comprehensive assessment of both substantive information quality and accountability standards. The use of non-parametric statistical methods is appropriate given the ordinal nature of the primary outcome measures, and the inclusion of inter-rater reliability assessment ensures the credibility of the quality ratings. The results of the analyses described herein are presented in Chapter 3.",
        first_line_indent=1.27,
        space_before=6
    )

    # ---- Save the document ----
    output_path = r"c:\Users\ayoth\Downloads\Masters Thesis\chapters\chapter2_methodology.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Chapter 2 saved successfully to: {output_path}")


if __name__ == "__main__":
    create_chapter2()
