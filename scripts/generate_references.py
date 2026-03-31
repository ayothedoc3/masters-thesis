"""
Generate the unified References section for the thesis.
Combines references from all chapter builders into a single APA 7 list.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Complete reference list -- all citations across all chapters, APA 7th edition
# Refreshed for the Chapter 1 age-compliance pass, 2026-03-31
REFERENCES = [
    "Afful-Dadzie, E., Afful-Dadzie, A., & Egala, S. B. (2023). Social media in health communication: A literature review of information quality. Health Information Management Journal, 52(1), 3-17. https://doi.org/10.1177/1833358321992683",
    "Azer, S. A. (2020). Are DISCERN and JAMA suitable instruments for assessing YouTube videos on thyroid cancer? Methodological concerns. Journal of Cancer Education, 35(6), 1267-1277. https://doi.org/10.1007/s13187-020-01763-9",
    "Balogun, B. A., Hogden, A., Kemp, N., Yang, L., & Agaliotis, M. (2023). Public health agencies' use of social media for communication during pandemics: A scoping review of the literature. Osong Public Health and Research Perspectives, 14(4), 235-251. https://doi.org/10.24171/j.phrp.2023.0095",
    "Baqraf, Y. K. A., Keikhosrokiani, P., & Al-Rawashdeh, M. (2023). Evaluating online health information quality using machine learning and deep learning: A systematic literature review. DIGITAL HEALTH, 9, 20552076231212296. https://doi.org/10.1177/20552076231212296",
    "Charnock, D., Shepperd, S., Needham, G., & Gann, R. (1999). DISCERN: An instrument for judging the quality of written consumer health information on treatment choices. Journal of Epidemiology and Community Health, 53(2), 105-111. https://doi.org/10.1136/jech.53.2.105",
    "Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.",
    "Cooke, A., Smith, D., & Booth, A. (2012). Beyond PICO: The SPIDER tool for qualitative evidence synthesis. Qualitative Health Research, 22(10), 1435-1443. https://doi.org/10.1177/1049732312452938",
    "Daraz, L., Dogu, C., Houde, V., Bouseh, S., & Morshed, K. G. (2024). Assessing credibility: Quality criteria for patients, caregivers, and the public in online health information-A qualitative study. Journal of Patient Experience, 11, 23743735241259440. https://doi.org/10.1177/23743735241259440",
    "Denniss, E., Lindberg, R., & McNaughton, S. A. (2023). Quality and accuracy of online nutrition-related information: A systematic review of content analysis studies. Public Health Nutrition, 26(7), 1345-1357. https://doi.org/10.1017/S1368980023000873",
    "Eysenbach, G., Powell, J., Kuss, O., & Sa, E. R. (2002). Empirical studies assessing the quality of health information for consumers on the World Wide Web: A systematic review. JAMA, 287(20), 2691-2700. https://doi.org/10.1001/jama.287.20.2691",
    "Feinstein, A. R., & Cicchetti, D. V. (1990). High agreement but low kappa: I. The problems of two paradoxes. Journal of Clinical Epidemiology, 43(6), 543-549. https://doi.org/10.1016/0895-4356(90)90158-L",
    "Fritz, C. O., Morris, P. E., & Richler, J. J. (2012). Effect size estimates: Current use, calculations, and interpretation. Journal of Experimental Psychology: General, 141(1), 2-18. https://doi.org/10.1037/a0024338",
    "Haghighi, R., & Farhadloo, M. (2025). Quality assessment of health information on social media during a public health crisis: Infodemiology study. JMIR Infodemiology, 5, e70756. https://doi.org/10.2196/70756",
    "Hua, Z., Song, Y., Liu, Q., & Chen, H. (2025). Factors influencing eHealth literacy worldwide: Systematic review and meta-analysis. Journal of Medical Internet Research, 27, e50313. https://doi.org/10.2196/50313",
    "Jia, X., Pang, Y., & Liu, L. S. (2021). Online health information seeking behavior: A systematic review. Healthcare, 9(12), 1740. https://doi.org/10.3390/healthcare9121740",
    "Jiang, X., Wang, L., Leng, Y., Xie, R., Li, C., Nie, Z., Liu, D., & Wang, G. (2024). The level of electronic health literacy among older adults: A systematic review and meta-analysis. Archives of Public Health, 82(1), 204. https://doi.org/10.1186/s13690-024-01428-9",
    "Kim, K., Shin, S., Kim, S., & Lee, E. (2023). The relation between eHealth literacy and health-related behaviors: Systematic review and meta-analysis. Journal of Medical Internet Research, 25, e40778. https://doi.org/10.2196/40778",
    "Kong, W., Song, S., Zhao, Y. C., Zhu, Q., & Sha, L. (2021). TikTok as a health information source: Assessment of the quality of information in diabetes-related videos. Journal of Medical Internet Research, 23(9), e30409. https://doi.org/10.2196/30409",
    "Koo, T. K., & Li, M. Y. (2016). A guideline of selecting and reporting intraclass correlation coefficients for reliability research. Journal of Chiropractic Medicine, 15(2), 155-163. https://doi.org/10.1016/j.jcm.2016.02.012",
    "Krippendorff, K. (2018). Content analysis: An introduction to its methodology (4th ed.). SAGE Publications.",
    "Landis, J. R., & Koch, G. G. (1977). The measurement of observer agreement for categorical data. Biometrics, 33(1), 159-174. https://doi.org/10.2307/2529310",
    "Li, B., Liu, X., Zhang, Y., Wang, Y., & Zheng, L. (2024). Quality assessment of health science-related short videos on TikTok: A scoping review. International Journal of Medical Informatics, 186, 105426. https://doi.org/10.1016/j.ijmedinf.2024.105426",
    "Li, R., & Ye, H. (2025). Wellness misinformation on social media: A systematic review using social cognitive theory. Health Communication, 1-16. https://doi.org/10.1080/10410236.2025.2555614",
    "Li, Y., Ouyang, H., Lin, G., Peng, Y., Yao, J., & Chen, Y. (2025). Evaluation of the measurement properties of online health information quality assessment tools: A systematic review. International Journal of Nursing Sciences, 12(2), 130-136. https://doi.org/10.1016/j.ijnss.2025.02.015",
    "Liu, X.-J., Valdez, D., Parker, M. A., Mai, A., & Walsh-Buhi, E. R. (2025). Quality of cancer-related information on new media (2014-2023): Systematic review and meta-analysis. Journal of Medical Internet Research, 27, e73185. https://doi.org/10.2196/73185",
    "Marocolo, M., Meireles, A., de Souza, H. L. R., Mota, G. R., Oranchuk, D. J., Arriel, R. A., & Leite, L. H. R. (2021). Is social media spreading misinformation on exercise and health in Brazil? International Journal of Environmental Research and Public Health, 18(22), 11914. https://doi.org/10.3390/ijerph182211914",
    "McKinney, W. (2010). Data structures for statistical computing in Python. In S. van der Walt & J. Millman (Eds.), Proceedings of the 9th Python in Science Conference (pp. 56-61). https://doi.org/10.25080/Majora-92bf1922-00a",
    "Mohamed, F., & Shoufan, A. (2024). Users' experience with health-related content on YouTube: An exploratory study. BMC Public Health, 24, 86. https://doi.org/10.1186/s12889-023-17585-5",
    "Mohamed, H., Kittle, E., Nour, N., Hamed, R., Feeney, K., Salsberg, J., & Kelly, D. (2024). An integrative systematic review on interventions to improve layperson's ability to identify trustworthy digital health information. PLOS Digital Health, 3(10), e0000638. https://doi.org/10.1371/journal.pdig.0000638",
    "Norman, C. D., & Skinner, H. A. (2006). eHEALS: The eHealth Literacy Scale. Journal of Medical Internet Research, 8(4), e27. https://doi.org/10.2196/jmir.8.4.e27",
    "Osman, W., Mohamed, F., Elhassan, M., & Shoufan, A. (2022). Is YouTube a reliable source of health-related information? A systematic review. BMC Medical Education, 22, 382. https://doi.org/10.1186/s12909-022-03446-z",
    "Pan, B., Hembrooke, H., Joachims, T., Lorigo, L., Gay, G., & Granka, L. (2007). In Google we trust: Users' decisions on rank, position, and relevance. Journal of Computer-Mediated Communication, 12(3), 801-823. https://doi.org/10.1111/j.1083-6101.2007.00351.x",
    "Paul, B., & Headley-Johnson, S. (2025). The impact of social media on health behaviors, a systematic review. Healthcare, 13(21), 2763. https://doi.org/10.3390/healthcare13212763",
    "Pew Research Center. (2021). Social media use in 2021. https://www.pewresearch.org/internet/2021/04/07/social-media-use-in-2021/",
    "Rodriguez-Rodriguez, A. M., Blanco-Diaz, M., de la Fuente-Costa, M., Hernandez-Sanchez, S., Escobio-Prieto, I., & Casa\u00f1a, J. (2022). Review of the quality of YouTube videos recommending exercises for the COVID-19 lockdown. International Journal of Environmental Research and Public Health, 19(13), 8016. https://doi.org/10.3390/ijerph19138016",
    "Rosenstock, I. M. (1974). Historical origins of the Health Belief Model. Health Education Monographs, 2(4), 328-335. https://doi.org/10.1177/109019817400200403",
    "Schober, P., Boer, C., & Schwarte, L. A. (2018). Correlation coefficients: Appropriate use and interpretation. Anesthesia & Analgesia, 126(5), 1763-1768. https://doi.org/10.1213/ANE.0000000000002864",
    "Silberg, W. M., Lundberg, G. D., & Musacchio, R. A. (1997). Assessing, controlling, and assuring the quality of medical information on the internet: Caveat lector et viewor-Let the reader and viewer beware. JAMA, 277(15), 1244-1245. https://doi.org/10.1001/jama.1997.03540390074039",
    "Singh, B., Ahmed, M., Staiano, A., Gough, C., & Vandelanotte, C. (2024). A systematic umbrella review and meta-meta-analysis of eHealth and mHealth interventions for improving lifestyle behaviours. npj Digital Medicine, 7(1), 179. https://doi.org/10.1038/s41746-024-01172-y",
    "Snelson, C. L. (2016). Qualitative and mixed methods social media research: A review of the literature. International Journal of Qualitative Methods, 15(1), 1-15. https://doi.org/10.1177/1609406915624574",
    "Stimpson, J. P., Park, S., Adhikari, E. H., Nelson, D. B., & Ortega, A. N. (2025). Perceived health misinformation on social media and public trust in health care. Medical Care, 63(9), 686-693. https://doi.org/10.1097/MLR.0000000000002180",
    "Suarez-Lledo, V., & Alvarez-Galvez, J. (2021). Prevalence of health misinformation on social media: Systematic review. Journal of Medical Internet Research, 23(1), e17187. https://doi.org/10.2196/17187",
    "Sullivan, G. M., & Artino, A. R. (2013). Analyzing and interpreting data from Likert-type scales. Journal of Graduate Medical Education, 5(4), 541-542. https://doi.org/10.4300/JGME-5-4-18",
    "Sun, F., Zheng, S., & Wu, J. (2023). Quality of information in gallstone disease videos on TikTok: Cross-sectional study. Journal of Medical Internet Research, 25, e39162. https://doi.org/10.2196/39162",
    "Thomas, D. D., Xu, L., Yu, B., Alanis, O., Adamek, J., Canton, I., Lin, X., Luo, Y., & Mullen, S. P. (2025). Physical activity misinformation on social media: Systematic review. JMIR Infodemiology, 5, e62760. https://doi.org/10.2196/62760",
    "Tomczak, M., & Tomczak, E. (2014). The need to report effect size estimates revisited: An overview of some recommended measures of effect size. Trends in Sport Sciences, 21(1), 19-25.",
    "Tripodi, F. (2018). Searching for alternative facts: Analyzing scriptural inference in conservative news practices. Data & Society Research Institute.",
    "Vallat, R. (2018). Pingouin: Statistics in Python. Journal of Open Source Software, 3(31), 1026. https://doi.org/10.21105/joss.01026",
    "Virtanen, P., Gommers, R., Oliphant, T. E., Haberland, M., Reddy, T., Cournapeau, D., Burovski, E., Peterson, P., Weckesser, W., Bright, J., van der Walt, S. J., Brett, M., Wilson, J., Millman, K. J., Mayorov, N., Nelson, A. R. J., Jones, E., Kern, R., Larson, E., ... SciPy 1.0 Contributors. (2020). SciPy 1.0: Fundamental algorithms for scientific computing in Python. Nature Methods, 17(3), 261-272. https://doi.org/10.1038/s41592-019-0686-2",
    "Wang, R. Y., & Strong, D. M. (1996). Beyond accuracy: What data quality means to data consumers. Journal of Management Information Systems, 12(4), 5-33. https://doi.org/10.1080/07421222.1996.11518099",
    "World Health Organization. (2020). WHO guidelines on physical activity and sedentary behaviour. World Health Organization. https://www.who.int/publications/i/item/9789240015128",
]


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    h = doc.add_heading("REFERENCES", level=1)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("")

    sorted_refs = sorted(REFERENCES, key=lambda x: x.lower())

    for ref in sorted_refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    path = os.path.join(ROOT, "references", "references.docx")
    doc.save(path)
    print(f"References saved to {path} ({len(sorted_refs)} entries)")


if __name__ == "__main__":
    main()
