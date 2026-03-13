"""
Convert Chapter 1 Literature Review from markdown to formatted .docx
with LSU formatting (Times New Roman 12pt, 1.5 spacing, 1" margins, justified).
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INPUT = os.path.join(ROOT, "chapters", "chapter1_literature_review.md")
OUTPUT = os.path.join(ROOT, "chapters", "chapter1_literature_review.docx")


def create_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Times New Roman"
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(14)
        elif level == 2:
            hs.font.size = Pt(13)
        else:
            hs.font.size = Pt(12)
        hs.paragraph_format.line_spacing = 1.5
        hs.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    return doc


def main():
    with open(INPUT, "r", encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = create_doc()
    in_references = False

    for line in lines:
        stripped = line.strip()

        if not stripped:
            if not in_references:
                doc.add_paragraph("")
            continue

        # Chapter title
        if stripped.startswith("CHAPTER 1"):
            h = doc.add_heading(stripped, level=1)
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # Section headings (1.1, 1.2, etc.)
        if re.match(r"^1\.\d\s", stripped):
            h = doc.add_heading(stripped, level=2)
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.color.rgb = RGBColor(0, 0, 0)
            continue

        # REFERENCES heading
        if stripped == "REFERENCES":
            in_references = True
            continue  # We'll handle references in the main assembly

        # Skip reference entries (they go in the unified reference list)
        if in_references:
            continue

        # Numbered items (research questions)
        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.left_indent = Inches(0.5)
            r = p.add_run(stripped)
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(stripped)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)

    doc.save(OUTPUT)
    print(f"Chapter 1 saved to {OUTPUT}")

    # Count paragraphs
    verify = Document(OUTPUT)
    count = sum(1 for p in verify.paragraphs if p.text.strip())
    print(f"Non-empty paragraphs: {count}")


if __name__ == "__main__":
    main()
