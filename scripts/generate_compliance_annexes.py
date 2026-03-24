"""
Generate additional annex documents required for LSU compliance:
- Draft scientific article / manuscript
- Conference evidence
- Consultations timesheet placeholder
- Ethics evidence placeholder
"""
from __future__ import annotations

from io import BytesIO
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
APPENDICES = os.path.join(ROOT, "appendices")
os.makedirs(APPENDICES, exist_ok=True)
CONFERENCE_CERT_PDF = os.path.join(os.path.dirname(ROOT), "conference certificate lsu.pdf")


def set_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_heading(doc: Document, text: str, *, center: bool = True, size: int = 14) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True


def add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_cm: float | None = None,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent_cm is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def generate_article_annex() -> None:
    doc = Document()
    set_style(doc)

    add_heading(doc, "APPENDIX E: DRAFT SCIENTIFIC ARTICLE PREPARED FOR JOURNAL SUBMISSION")
    add_para(doc, "Target journal format: JMIR Public Health and Surveillance (manuscript-style draft).", italic=True)
    add_para(
        doc,
        "Quality of Physical Activity Information on Traditional Websites Versus Social Media Platforms: "
        "A Cross-Sectional Content Analysis Using DISCERN and JAMA Benchmarks",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_para(doc, "Ayokunle Ademola-John", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Lithuanian Sports University, Kaunas, Lithuania", align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "Abstract", center=False, size=12)
    add_para(
        doc,
        "Background: Online physical activity information is increasingly consumed through both traditional "
        "websites and social media platforms, yet direct cross-platform comparisons remain limited. "
        "Objective: To compare the quality of physical activity information on traditional websites, "
        "YouTube, TikTok, and Instagram using DISCERN and JAMA benchmarks. Methods: A cross-sectional "
        "content analysis was conducted on 197 publicly accessible English-language items retrieved with "
        "five standardized search terms. Quality was assessed with the 16-item DISCERN instrument and four "
        "JAMA accountability criteria. Inter-rater reliability was evaluated on 40 items. Results: Quality "
        "differed significantly across platforms, with websites performing best, followed by YouTube, "
        "TikTok, and Instagram. JAMA compliance was low across all platforms. Creator type was associated "
        "with quality, and YouTube views and likes showed negative correlations with DISCERN scores. "
        "Conclusions: Traditional websites remain the most reliable source of physical activity information, "
        "whereas short-form social media content demonstrates weaker quality and transparency.",
        first_line_indent_cm=1.25,
    )
    add_para(doc, "Keywords: physical activity; information quality; social media; DISCERN; JAMA", italic=True)

    add_heading(doc, "Introduction", center=False, size=12)
    add_para(
        doc,
        "The public now accesses physical activity guidance through a fragmented digital ecosystem in which "
        "institutional websites compete with user-generated and algorithmically promoted social media content. "
        "Although this improves reach, it also raises concerns about the credibility, transparency, and safety "
        "of the information users encounter. Prior studies have often examined one platform at a time, making "
        "it difficult to compare the quality of information across major online environments. This study "
        "addressed that gap by evaluating four platform types under a common protocol.",
        first_line_indent_cm=1.25,
    )

    add_heading(doc, "Methods", center=False, size=12)
    add_para(
        doc,
        "A cross-sectional content analysis was performed using five standardized search terms: how to start "
        "exercising, best exercises to lose weight, strength training for beginners, physical activity "
        "guidelines, and home workout routine. The final audited sample comprised 197 items: 50 traditional "
        "websites, 50 YouTube videos, 47 TikTok posts, and 50 Instagram posts. Information quality was rated "
        "with the DISCERN instrument and the four JAMA criteria of authorship, attribution, disclosure, and "
        "currency. A second rater independently scored 40 items to assess reliability. Non-parametric tests "
        "were used because the main outcome measures were ordinal and non-normally distributed.",
        first_line_indent_cm=1.25,
    )

    add_heading(doc, "Results", center=False, size=12)
    add_para(
        doc,
        "A clear quality gradient emerged across platforms. Traditional websites achieved the highest DISCERN "
        "scores, followed by YouTube, TikTok, and Instagram. JAMA compliance was modest overall, with currency "
        "being the most frequently satisfied criterion and attribution and disclosure the weakest. Creator type "
        "was associated with quality, with healthcare professionals and certified fitness professionals "
        "outperforming general users. Engagement metrics did not consistently signal higher quality. On "
        "YouTube, higher views and likes were associated with lower DISCERN scores.",
        first_line_indent_cm=1.25,
    )

    add_heading(doc, "Discussion", center=False, size=12)
    add_para(
        doc,
        "The findings suggest that platform architecture matters for information quality. Traditional websites "
        "retain structural advantages, including greater space for context, attribution, and update history. "
        "Short-form social media, by contrast, appears more vulnerable to brevity, limited source disclosure, "
        "and engagement-driven ranking. The significant negative within-platform correlations observed on "
        "YouTube indicate that popularity should not be treated as a proxy for quality. These results support "
        "the need for stronger digital health literacy and platform-level quality safeguards.",
        first_line_indent_cm=1.25,
    )

    add_heading(doc, "Conclusions", center=False, size=12)
    add_para(
        doc,
        "Physical activity information quality differs substantially across digital platforms. Traditional "
        "websites remain the most reliable source, whereas short-form social media content is more likely to "
        "lack transparency and to provide lower-quality information. Professionals produce stronger content on "
        "average, but platform effects remain important. Future interventions should focus on content quality "
        "standards, source transparency, and platform-aware health communication strategies.",
        first_line_indent_cm=1.25,
    )

    path = os.path.join(APPENDICES, "appendix_e_article.docx")
    doc.save(path)
    print(f"Appendix E saved to {path}")


def generate_conference_annex() -> None:
    doc = Document()
    set_style(doc)
    add_heading(doc, "APPENDIX F: EVIDENCE OF SCIENTIFIC CONFERENCE PRESENTATION")

    if os.path.exists(CONFERENCE_CERT_PDF) and fitz is not None:
        add_para(
            doc,
            "Official conference certificate confirming presentation of the thesis findings, "
            "included in accordance with LSU regulation clause 41.2.",
            first_line_indent_cm=1.25,
        )
        pdf = fitz.open(CONFERENCE_CERT_PDF)
        try:
            page = pdf.load_page(0)
            pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_stream = BytesIO(pix.tobytes("png"))
        finally:
            pdf.close()

        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_stream, width=Inches(6.2))
    else:
        add_para(
            doc,
            "This annex is reserved for the conference certificate, conference program, or other official "
            "document confirming presentation of the thesis findings at a scientific conference, as required "
            "by LSU regulation clause 41.2.",
            first_line_indent_cm=1.25,
        )
        add_para(doc, "")
        add_para(doc, "Insert the signed or official document in this annex before submission.", bold=True)

    path = os.path.join(APPENDICES, "appendix_f_conference_evidence.docx")
    doc.save(path)
    print(f"Appendix F saved to {path}")


def generate_document_placeholder(filename: str, title: str, body: str) -> None:
    doc = Document()
    set_style(doc)
    add_heading(doc, title)
    add_para(doc, body, first_line_indent_cm=1.25)
    add_para(doc, "")
    add_para(doc, "Insert the signed or official document in this annex before submission.", bold=True)
    path = os.path.join(APPENDICES, filename)
    doc.save(path)
    print(f"{title} saved to {path}")


if __name__ == "__main__":
    generate_article_annex()
    generate_conference_annex()
    generate_document_placeholder(
        "appendix_g_consultations_timesheet.docx",
        "APPENDIX G: CONSULTATIONS TIMESHEET",
        "This annex is reserved for the completed consultations timesheet coordinated with the scientific "
        "supervisor, as required by LSU regulation clause 12.2.",
    )
    generate_document_placeholder(
        "appendix_h_ethics_evidence.docx",
        "APPENDIX H: ETHICS PERMISSION OR FORMAL EXEMPTION",
        "This annex is reserved for the official ethics approval, ethics exemption, or Study Programme "
        "Committee decision documenting why formal ethics approval was not required, as referenced in LSU "
        "regulation clauses 14.2 and 41.3.",
    )
    print("\nAdditional compliance annexes created successfully.")
