#!/usr/bin/env python3
"""
Generate Chapter 4: Discussion for Master's Thesis
Lithuanian Sports University
Author: Ayokunle Ademola-John
"""

import subprocess
import sys

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def format_run(run, bold=False, italic=False, size=12, font_name="Times New Roman", color=None):
    """Format a run with specified properties."""
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>'
    )
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing


def add_formatted_paragraph(doc, text, bold=False, italic=False, size=12,
                             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=0, space_after=6, line_spacing=1.5,
                             first_line_indent=None):
    """Add a fully formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    format_run(run, bold=bold, italic=italic, size=size)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_heading_custom(doc, text, level=1, size=14, space_before=12, space_after=6):
    """Add a custom heading with Times New Roman formatting."""
    p = doc.add_paragraph(style=f"Heading {1 if level == 1 else 2 if level == 2 else 3}")
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    format_run(run, bold=True, size=size)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=1.5)
    return p


def add_mixed_paragraph(doc, parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         space_before=0, space_after=6, line_spacing=1.5,
                         first_line_indent=None):
    """
    Add a paragraph with mixed formatting.
    parts: list of tuples (text, bold, italic)
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    for text, bold, italic in parts:
        run = p.add_run(text)
        format_run(run, bold=bold, italic=italic, size=12)
    add_paragraph_spacing(p, before=space_before, after=space_after, line_spacing=line_spacing)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


# ============================================================
# DOCUMENT SETUP
# ============================================================

doc = Document()

# Page setup: A4, 1-inch margins
for section in doc.sections:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
rPr = style.element.get_or_add_rPr()
rFonts = parse_xml(
    f'<w:rFonts {nsdecls("w")} w:ascii="Times New Roman" '
    f'w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
)
rPr.insert(0, rFonts)

# ============================================================
# CHAPTER 4: CONSIDERATIONS
# ============================================================

add_heading_custom(doc, "CHAPTER 4. CONSIDERATIONS", level=1, size=14, space_before=0, space_after=12)

# Introductory paragraph
add_formatted_paragraph(
    doc,
    "This chapter interprets the principal findings of the cross-sectional content analysis "
    "in light of existing literature, evaluates the study's methodological strengths and "
    "limitations, and outlines implications for public health practice and future research. "
    "The overarching aim was to determine whether the quality of physical activity (PA) "
    "information differs across traditional websites, YouTube, TikTok, and Instagram when "
    "assessed with the DISCERN instrument and JAMA benchmarks, and whether creator type "
    "and audience engagement moderate that quality.",
    first_line_indent=1.25
)

# ============================================================
# 4.1 Summary of Key Findings
# ============================================================

add_heading_custom(doc, "4.1 Summary of Key Findings", level=2, size=12, space_before=12, space_after=6)

add_formatted_paragraph(
    doc,
    "Three principal findings emerged from the analysis of 197 audited items across four platforms. "
    "First, a clear and statistically significant quality gradient was observed: traditional "
    "websites produced the highest DISCERN scores (Mdn = 48, classified as fair-to-good), "
    "followed by YouTube (Mdn = 39, fair-to-poor), TikTok (Mdn = 31, poor), and Instagram "
    "(Mdn = 21, very poor). The Kruskal-Wallis test confirmed this pattern, H(3) = 136.58, "
    "p < .001, with a large effect size (\u03b7\u00b2 = .68), and all six pairwise comparisons reached "
    "statistical significance after Bonferroni correction. This gradient suggests that "
    "platform architecture\u2014particularly the space available for textual detail, referencing, "
    "and authorship disclosure\u2014plays a decisive role in the completeness and reliability "
    "of health content.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Second, creator type significantly moderated DISCERN scores, with healthcare and "
    "certified fitness professionals scoring higher (Mdn = 39.5) than non-professional "
    "creators (Mdn = 35.0; Mann-Whitney U, p = .033). However, the effect was small "
    "(r = \u2212.204), indicating that professional credentials alone do not guarantee "
    "high-quality output. Notably, organisations dominated the traditional website space "
    "(78% of website items), whereas fitness influencers and general users were the "
    "predominant creators on YouTube and TikTok. This distributional imbalance partly "
    "explains the platform quality gradient, as the platforms with the lowest scores were "
    "also those with the fewest professionally credentialed authors.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Third, engagement metrics bore virtually no relationship to information quality. "
    "The overall Spearman correlation between views and DISCERN total was negligible "
    "(\u03c1 = .017). Platform-specific analyses revealed a negative correlation on YouTube "
    "(higher view counts associated with lower quality) and, unexpectedly, a positive "
    "correlation on TikTok. These divergent patterns suggest that algorithmic "
    "recommendation systems reward different content attributes on each platform, none "
    "of which appear to be information accuracy or completeness.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Regarding JAMA benchmark compliance, the overall JAMA Total score did not differ "
    "significantly across platforms (p = .143). However, analysis of individual JAMA "
    "criteria revealed significant differences for Attribution (\u03c7\u00b2, p < .001), "
    "Disclosure (p = .001), and Currency (p < .001). Social media posts almost universally "
    "met the Currency criterion because post dates are inherently visible, yet they "
    "performed poorly on Attribution and Disclosure. This pattern highlights a structural "
    "transparency gap: social media platforms do not incentivise or facilitate source "
    "citation or conflict-of-interest disclosure in the way that editorial standards do "
    "for traditional web publishing.",
    first_line_indent=1.25
)

# ============================================================
# 4.2 Comparison with Existing Literature
# ============================================================

add_heading_custom(doc, "4.2 Comparison with Existing Literature", level=2, size=12, space_before=12, space_after=6)

add_formatted_paragraph(
    doc,
    "The quality gradient observed in this study is broadly consistent with earlier work "
    "on online health information. The overall mean DISCERN score of 38.68 out of 80 aligns "
    "closely with the pooled estimate from the 2025 JMIR meta-analysis, which reported a "
    "DISCERN score of approximately 43.58 out of 100 (equivalent to roughly 34.9 on an "
    "80-point scale) across diverse health topics (Liu et al., 2025). That our sample slightly "
    "exceeded this benchmark is encouraging, but it also underscores that even the best-"
    "performing platform\u2014traditional websites\u2014only reached the lower boundary of the "
    "\u201cgood\u201d category. This finding resonates with Eysenbach's (2002) foundational observation "
    "that health information quality on the internet is highly variable, a conclusion that "
    "remains valid more than two decades later.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "The particularly low scores for Instagram and TikTok mirror findings from recent "
    "platform-specific studies. Li et al. (2024) concluded that most health-related TikTok "
    "videos assessed to date are of moderate-to-low quality and that medically credentialed "
    "creators typically account for only a minority of content producers. In the present "
    "study, healthcare professionals accounted for 16% of TikTok items, a proportion that "
    "was still insufficient to lift TikTok's overall quality above the \u201cpoor\u201d threshold. "
    "On JAMA compliance, Haghighi "
    "and Farhadloo (2025) found that 81% of health-related Twitter content was of low "
    "quality and 95% met two or fewer of the four JAMA criteria. Our social media data "
    "follow the same pattern, with the majority of TikTok and Instagram items failing to "
    "meet Attribution and Disclosure criteria.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "The negligible overall correlation between engagement and quality is consistent with "
    "review-level evidence that surface popularity metrics are poor proxies for the quality "
    "of health information on YouTube and other social platforms (Osman et al., 2022; "
    "Mohamed & Shoufan, 2024). Our observation that YouTube exhibited significant negative "
    "views-quality and likes-quality relationships supports the hypothesis that "
    "algorithmically promoted content on video platforms may prioritise entertainment value "
    "and emotional appeal over informational rigour (Moorhead et al., 2013). In contrast, "
    "TikTok and Instagram showed no significant within-platform correlations, reinforcing "
    "the interpretation that aggregate engagement-quality associations are largely a "
    "between-platform artefact rather than evidence that highly engaged content is higher "
    "quality.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "The finding that creator type moderates quality\u2014albeit with a small effect\u2014parallels "
    "evidence from the broader health literacy literature. Norman and Skinner (2006) argued "
    "that both provider-side and consumer-side competencies shape the quality of digital "
    "health exchanges. The present data suggest that provider-side credentialing is a "
    "necessary but not sufficient condition for high-quality content: professional creators "
    "still frequently omitted references and disclosure statements, particularly when "
    "operating within the formatting constraints of social media posts. This implies that "
    "platform design, not just author expertise, constrains the extent to which "
    "accountability criteria can be fulfilled.",
    first_line_indent=1.25
)

# ============================================================
# 4.3 Implications for Public Health Practice
# ============================================================

add_heading_custom(doc, "4.3 Implications for Public Health Practice", level=2, size=12, space_before=12, space_after=6)

add_formatted_paragraph(
    doc,
    "These findings carry several implications for public health practitioners and policy "
    "makers. First, the clear platform quality gradient argues against treating the internet "
    "as a monolithic information source. Public health messaging strategies should be "
    "calibrated to the specific affordances and limitations of each platform. For instance, "
    "short-form video platforms such as TikTok and Instagram Reels impose severe constraints "
    "on citation, attribution, and nuanced explanation; public health organisations "
    "seeking to use these platforms must develop content formats that maximise accuracy "
    "within those constraints, such as pinned comments with references or standardised "
    "disclosure overlays.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Second, the absence of a positive engagement-quality relationship suggests that "
    "high-quality health information does not organically reach large audiences on social "
    "media. Public health agencies may therefore need to invest in paid promotion or "
    "influencer partnerships to ensure that evidence-based PA information achieves "
    "competitive visibility. Broader public-health communication and infodemic-management "
    "literature already supports strategic use of digital platforms alongside stronger "
    "quality assurance mechanisms (Swire-Thompson & Lazer, 2020), and the present data "
    "reinforce that need for physical activity information specifically.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Third, the dominance of non-professional creators on social media platforms calls for "
    "educational interventions targeting both content creators and consumers. Creator-side "
    "interventions could include platform-integrated prompts to cite sources or disclose "
    "conflicts of interest, while consumer-side interventions should focus on enhancing "
    "eHealth literacy (Norman & Skinner, 2006) so that users can critically appraise "
    "PA content regardless of the platform on which it is encountered. Fox and Duggan (2013) "
    "documented that a substantial proportion of health information seekers do not verify "
    "the sources they consult; the present data suggest that this problem is compounded on "
    "platforms where source verification is structurally difficult.",
    first_line_indent=1.25
)

# ============================================================
# 4.4 Strengths of the Study
# ============================================================

add_heading_custom(doc, "4.4 Strengths of the Study", level=2, size=12, space_before=12, space_after=6)

add_formatted_paragraph(
    doc,
    "This study offers several methodological strengths. To the best of our knowledge, it "
    "is the first investigation to compare all four platform types\u2014traditional websites, "
    "YouTube, TikTok, and Instagram\u2014within a single study using both the DISCERN instrument "
    "and JAMA benchmarks. Previous research has typically examined one or two platforms in "
    "isolation, limiting cross-platform inference. The present design enables direct, "
    "standardised comparison across the platforms most commonly accessed for PA information.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "The final audited sample of 197 items (50 websites, 50 YouTube videos, 47 TikTok "
    "videos, and 50 Instagram posts) is larger than many content analysis studies in this "
    "domain and was drawn using a systematic protocol of five representative search terms "
    "applied uniformly across all platforms. Inter-rater reliability was excellent, "
    "with an intraclass correlation coefficient (ICC) of 0.932 for DISCERN scores, well "
    "above the commonly accepted threshold of 0.75 (Koo & Li, 2016). This level of "
    "agreement strengthens confidence that the observed quality differences reflect genuine "
    "content variation rather than rater inconsistency.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Additionally, the use of both DISCERN (measuring reliability and treatment information "
    "quality) and JAMA benchmarks (measuring accountability through authorship, attribution, "
    "disclosure, and currency) provides a more comprehensive quality profile than either "
    "instrument alone. This dual-instrument approach is especially valuable because DISCERN "
    "and JAMA capture partially overlapping but distinct dimensions of information quality "
    "(Charnock et al., 1999; Silberg et al., 1997).",
    first_line_indent=1.25
)

# ============================================================
# 4.5 Limitations
# ============================================================

add_heading_custom(doc, "4.5 Limitations", level=2, size=12, space_before=12, space_after=6)

add_formatted_paragraph(
    doc,
    "Several limitations should be considered when interpreting these results. First, the "
    "cross-sectional design captures a single snapshot of platform content. Online health "
    "information is dynamic; content is continuously created, updated, and removed. The "
    "quality landscape observed at the time of data collection may not generalise to other "
    "time periods, and longitudinal designs would be needed to track quality trends.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Second, the study was limited to English-language content. Physical activity "
    "information in other languages may differ in quality and creator composition, and "
    "findings may not transfer to non-English-speaking populations. Third, the scoring of "
    "video-based content on YouTube, TikTok, and Instagram relied on titles, descriptions, "
    "captions, and visible on-screen text rather than full audiovisual analysis of spoken "
    "content. This approach was necessary for feasibility and consistency with the text-"
    "oriented DISCERN instrument, but it may underestimate the informational depth of "
    "videos whose quality resides primarily in verbal explanation rather than written "
    "metadata.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Fourth, the sampling strategy\u2014selecting the top 10 results per search term\u2014"
    "captures the most visible content but may not represent the full breadth of PA "
    "information available on each platform. Users who scroll beyond the first page or "
    "receive personalised algorithmic recommendations may encounter content of different "
    "quality. Fifth, although DISCERN is the most widely validated tool for assessing "
    "consumer health information, it was originally designed for written patient information "
    "leaflets (Charnock et al., 1999). Applying it to social media posts required "
    "operationalisation decisions (detailed in Chapter 2) that, while transparent and "
    "systematic, represent an adaptation beyond the instrument\u2019s original scope.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Sixth, all searches were conducted from a single geographic location, meaning that "
    "search engine and platform algorithms may have returned location-specific results. "
    "Users in other regions could encounter a different set of top results for the same "
    "queries. Finally, engagement metrics were not uniformly available across all platforms; "
    "in particular, Instagram view counts were not captured for all items, which limits the "
    "strength of engagement-quality correlations for that platform.",
    first_line_indent=1.25
)

# ============================================================
# 4.6 Recommendations for Future Research
# ============================================================

add_heading_custom(doc, "4.6 Recommendations for Future Research", level=2, size=12, space_before=12, space_after=6)

add_formatted_paragraph(
    doc,
    "Future research should address the limitations identified above while extending the "
    "scope of cross-platform quality assessment. First, longitudinal studies are needed to "
    "determine whether the observed quality gradient is stable over time or whether it "
    "shifts as platforms evolve their content policies and recommendation algorithms. "
    "Repeated assessments at regular intervals (e.g., every six months) would also reveal "
    "whether public health interventions targeting content quality produce measurable "
    "improvements.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Second, studies should incorporate multilingual samples and diverse geographic search "
    "contexts to assess the generalisability of the present findings. The quality of PA "
    "information may vary substantially across languages and health systems, and cross-"
    "cultural comparisons would inform the design of context-appropriate interventions.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Third, the development of quality assessment instruments specifically designed for "
    "audiovisual and short-form content would improve the validity of social media "
    "evaluations. While DISCERN remains the gold standard for written health information, "
    "its application to 15-second TikTok videos or image-only Instagram posts stretches the "
    "instrument beyond its intended use. A validated tool that accounts for the multimodal "
    "nature of social media content\u2014including spoken narration, visual demonstrations, "
    "and on-screen text\u2014would be a significant methodological advance.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Fourth, experimental and quasi-experimental designs should examine the causal impact "
    "of information quality on health behaviour. The present study establishes that quality "
    "varies across platforms, but it cannot determine whether exposure to low-quality PA "
    "information leads to suboptimal exercise behaviours or, conversely, whether high-"
    "quality content promotes adherence to PA guidelines. Linking content quality data to "
    "user behaviour outcomes would strengthen the evidence base for quality-improvement "
    "interventions.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Fifth, research should investigate the role of platform design features in shaping "
    "content quality. A/B testing of interface modifications\u2014such as mandatory source-"
    "citation fields, credential-verification badges, or algorithmic boosts for referenced "
    "content\u2014could identify practical mechanisms for improving health information quality "
    "at scale. Such studies would require collaboration with platform operators but could "
    "yield high-impact, population-level benefits.",
    first_line_indent=1.25
)

add_formatted_paragraph(
    doc,
    "Finally, the relationship between engagement metrics and quality deserves deeper "
    "investigation using more granular data. Access to algorithmic ranking signals, watch "
    "time, completion rates, and share-to-view ratios\u2014data that are not publicly available "
    "but may be accessible through platform research APIs\u2014would enable more nuanced "
    "modelling of how recommendation systems interact with information quality. "
    "Understanding these dynamics is essential for designing interventions that align "
    "algorithmic incentives with public health objectives.",
    first_line_indent=1.25
)

# ============================================================
# SAVE DOCUMENT
# ============================================================

output_dir = r"c:\Users\ayoth\Downloads\Masters Thesis\chapters"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "chapter4_discussion.docx")
doc.save(output_path)
print(f"Chapter 4 saved to: {output_path}")
print("Done.")
